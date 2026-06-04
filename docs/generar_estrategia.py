from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Estilos globales ──────────────────────────────────────────────────────────
ORANGE  = RGBColor(0xFF, 0x9B, 0x11)
DARK    = RGBColor(0x2E, 0x2E, 0x2E)
GRAY    = RGBColor(0x50, 0x50, 0x50)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
LGRAY   = RGBColor(0xF5, 0xF5, 0xF5)

# Márgenes
sections = doc.sections
for s in sections:
    s.top_margin    = Cm(2)
    s.bottom_margin = Cm(2)
    s.left_margin   = Cm(2.5)
    s.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_cell_bg(cell, color_hex):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  color_hex)
    tcPr.append(shd)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size  = Pt(18)
        run.font.color.rgb = ORANGE
    elif level == 2:
        run.font.size  = Pt(14)
        run.font.color.rgb = DARK
    elif level == 3:
        run.font.size  = Pt(12)
        run.font.color.rgb = GRAY
    else:
        run.font.size  = Pt(11)
        run.font.color.rgb = ORANGE
    return p

def add_body(doc, text, bold=False, italic=False, size=10):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_bullet(doc, text, size=10):
    p   = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(2)
    return p

def make_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    # Header row
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        c   = hrow.cells[i]
        set_cell_bg(c, 'FF9B11')
        run = c.paragraphs[0].add_run(h)
        run.bold            = True
        run.font.color.rgb  = WHITE
        run.font.size       = Pt(9)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c   = t.rows[ri + 1].cells[ci]
            run = c.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
            if ri % 2 == 0:
                set_cell_bg(c, 'F5F5F5')
    # Column widths
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[ci].width = Cm(w)
    doc.add_paragraph()
    return t

def page_break(doc):
    doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# PORTADA
# ═══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Cm(4)
r = p.add_run('ESTRATEGIA GOOGLE ADS 2025–2026')
r.bold = True; r.font.size = Pt(24); r.font.color.rgb = ORANGE
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
r = p.add_run('Sistemas y Soluciones Digitales SRL · HUENU')
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = DARK
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Venta, alquiler y distribución de equipos de impresión digital\n'
              'Av. Belgrano 748, Piso 4, Of. 43 — CABA, Argentina\n'
              '(011) 4342-5742 · (911) 2787-0446 · info@sistemasysoluciones.com')
r.font.size = Pt(11); r.font.color.rgb = GRAY
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Documento confidencial — Junio 2026')
r.italic = True; r.font.size = Pt(10); r.font.color.rgb = GRAY
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# ÍNDICE
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, 'ÍNDICE', 1)
indice = [
    ('Fase 1', 'Auditoría del sitio web'),
    ('Fase 2', 'Estructura de campañas — 27 campañas'),
    ('Fase 3', 'Avisos RSA y extensiones por campaña'),
    ('Fase 4', 'Pujas, presupuesto y checklist de activación'),
]
for num, titulo in indice:
    p   = doc.add_paragraph()
    run = p.add_run(f'{num}  —  {titulo}')
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(3)

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# FASE 1 — AUDITORÍA
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, 'FASE 1 — AUDITORÍA DEL SITIO WEB', 1)
add_body(doc, 'Sitio auditado: www.sistemasysoluciones.com', bold=True)
add_body(doc, 'Fecha: Junio 2026')
doc.add_paragraph()

add_heading(doc, 'Observaciones generales', 2)
obs = [
    'CTAs disponibles en todo el sitio: botón +INFORMACIÓN por producto (abre formulario), WhatsApp (911) 2787-0446, teléfono (011) 4342-5742, email.',
    'No hay e-commerce ni carrito de compras.',
    'Las páginas de insumos (tintas, films, polvos, repuestos) son categorías sin catálogo individual de SKUs.',
    'Las landing pages de campaña están en servidor local pendientes de publicación.',
    'Los formularios de las 4 landing pages están conectados a WhatsApp mediante JavaScript (sin backend).',
    'www.huenu.ar tiene publicado únicamente el modelo HR1804UV (sitio ~40% completo).',
]
for o in obs:
    add_bullet(doc, o)
doc.add_paragraph()

add_heading(doc, 'Inventario — Equipos de Impresión Gran Formato (HUENU)', 2)
make_table(doc,
    ['Modelo', 'Tecnología', 'Ancho', 'Nivel', 'Estado'],
    [
        ('HR3002DT / HR6002DT / HR6004DT / HR6007DT', 'DTF Textil', '330–620 mm', 'Entry–Alto', 'Activa'),
        ('HR3001DU / HR6004DU', 'DTF UV + laminado', '320–620 mm', 'Entry–Mediano', 'Activa'),
        ('HR1602ES / HR1804ES / HR3204ES', 'Ecosolvente', '1600–3200 mm', 'Mediano–Alto', 'Activa'),
        ('HR1602SU / HR1804SU / HR1808SU / HR1902SU / HR2202SU', 'Sublimación', '1600–2220 mm', 'Mediano–Alta prod.', 'Activa'),
        ('HR1602UV / HR1804UV / HR2008UV / HR3204UV / HR3208UV / HR3212UV', 'UV Rollo', '1600–3200 mm', 'Mediano–Alta prod.', 'Activa'),
        ('HI1804UV', 'UV Híbrida', '1800 mm', 'Mediano', 'Activa'),
        ('HF3040 / HF6090', 'UV Cama Plana', '300×420 – 600×900 mm', 'Entry–Mediano', 'Activa'),
        ('HR3204ASV / HR3204BSV / HR3208ASV / HR3208BSV / HR3208CSV', 'Solvente', '3200 mm', 'Alta producción', 'Activa'),
    ],
    col_widths=[5.5, 3.5, 3, 3, 2]
)

add_heading(doc, 'Inventario — Equipos Alta Producción (Dlican, Sinajet, RISO)', 2)
make_table(doc,
    ['Modelo', 'Marca', 'Tecnología', 'Nivel', 'Geo'],
    [
        ('DLI-9060 a DLI-3325 (9 modelos)', 'Dlican', 'UV Cama Plana', 'Alta producción', 'Todo Argentina'),
        ('DLI-1688 a DLI-6600 (6 modelos)', 'Dlican', 'UV Híbrida', 'Alta producción', 'Todo Argentina'),
        ('Serie DG, DF-MT, DH, DF', 'Sinajet', 'Mesas de Corte CNC', 'Mediano–Alta prod.', 'Todo Argentina'),
        ('ComColor GD7330, FW5230', 'RISO', 'Inkjet alta velocidad', 'Alta producción', 'Todo Argentina'),
    ],
    col_widths=[4.5, 2.5, 3.5, 3, 3]
)

add_heading(doc, 'Inventario — Konica Minolta', 2)
make_table(doc,
    ['Modelo', 'Tipo', 'Velocidad', 'Segmento', 'Geo'],
    [
        ('Bizhub 558 / 658 / 958', 'Láser B/N', '55–95 ppm', 'Pequeño–Mediano', 'AMBA'),
        ('Bizhub C558 / C658', 'Láser Color', '55–65 ppm', 'Pequeño–Mediano', 'AMBA'),
        ('AccurioPress C3070/80 / C6100 / Bizhub Pro1100', 'Alta producción', '71–100 ppm', 'Industrial', 'AMBA'),
        ('Todos los modelos — modalidad alquiler', 'Alquiler', '—', 'Empresas e imprentas', 'AMBA'),
    ],
    col_widths=[4.5, 3, 2.5, 3, 2]
)

add_heading(doc, 'Inventario — Corte, Etiquetas e Insumos', 2)
make_table(doc,
    ['Marca / Producto', 'Tipo', 'Nivel', 'Geo'],
    [
        ('Tenneth FC5070E / FC7090A/U / FC9012 / FC1313U', 'Mesas de corte digital', 'Entry', 'AMBA'),
        ('Harpy R440 / Bizpress 13R', 'Impresoras de etiquetas', 'Mediano', 'AMBA'),
        ('Teneth RN3 / Duoblade WS Max / SX / Petit Pro', 'Cortadoras de etiquetas', 'Mediano', 'AMBA'),
        ('Tintas UV / Sub / Eco / DTF — Films DTF — Polvos DTF', 'Insumos gran formato', '—', 'Todo Argentina'),
        ('Tóners y repuestos KM originales', 'Insumos KM', '—', 'Todo Argentina'),
    ],
    col_widths=[5.5, 3.5, 2, 2]
)

add_heading(doc, 'Landing Pages disponibles', 2)
make_table(doc,
    ['Archivo', 'Producto', 'Formulario', 'Estado'],
    [
        ('lp-km-alquiler.html', 'KM Alquiler (Bizhub + AccurioPress)', 'WhatsApp (conectado)', 'Lista para publicar'),
        ('lp-hf6090.html', 'HUENU HF-6090 UV Cama Plana', 'WhatsApp (conectado)', 'Lista para publicar'),
        ('lp-hr1804uv.html', 'HUENU HR1804UV UV Rollo', 'WhatsApp (conectado)', 'Lista para publicar'),
        ('lp-teneth7090.html', 'Tenneth FC7090 Mesa de Corte', 'WhatsApp (conectado)', 'Lista para publicar'),
        ('lp-insumos-gran-formato.html', 'Insumos Gran Formato (tintas UV/Eco/Sub, DTF)', 'WhatsApp (conectado)', 'Lista para publicar'),
        ('lp-insumos-km.html', 'Insumos KM (tóners y repuestos Konica Minolta)', 'WhatsApp (conectado)', 'Lista para publicar'),
        ('index.html', 'Hub interno (índice de LPs)', '—', 'No usar para tráfico pago'),
    ],
    col_widths=[4, 5, 3.5, 3.5]
)

add_heading(doc, 'Brechas identificadas', 2)
brechas = [
    'No existe landing dedicada para HUENU Captación Distribuidores — crear en huenu.ar antes de activar campaña 25.',
    'Tóners KM no tienen página de producto individual — usar categoría /insumos-y-repuestos/ como destino.',
    'huenu.ar tiene solo HR1804UV publicado — resto de campañas HUENU interior apuntan a sistemasysoluciones.com.',
    'URL /gran-formato/uv-híbrido requiere tilde en la "i" (uv-h%C3%ADbrido) para no devolver 404.',
    'Mesas de corte Tenneth: URL correcta usa guion bajo: /mesas_de_corte/equipos-compatos-corte.',
]
for b in brechas:
    add_bullet(doc, b)

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# FASE 2 — ESTRUCTURA DE CAMPAÑAS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, 'FASE 2 — ESTRUCTURA DE CAMPAÑAS', 1)
add_body(doc, 'Total: 27 campañas · 1 cuenta Google Ads · 2 dominios (sistemasysoluciones.com · huenu.ar)', bold=True)
doc.add_paragraph()

add_heading(doc, 'Decisión de estructura — HUENU Usuario Final Argentina sin AMBA', 2)
add_body(doc, 'Se adoptaron 6 campañas por línea de producto (UV Rollo, UV Cama Plana, UV Híbrida, DTF Textil, DTF UV, Sublimación) + 1 campaña de Captación Distribuidores.')
add_body(doc, 'Justificación: cada línea tiene buyer persona y search intent diferenciados. Agrupar por "textil + eco" reduciría la relevancia del mensaje y la calidad de las landings de destino.')
doc.add_paragraph()

add_heading(doc, 'Resumen ejecutivo — 27 campañas', 2)
make_table(doc,
    ['#', 'Campaña', 'Marca(s)', 'Geo', 'Prioridad'],
    [
        ('01', 'KM Alquiler', 'Konica Minolta', 'AMBA', 'P1'),
        ('02', 'KM Venta Pequeños/Medios', 'Konica Minolta', 'AMBA', 'P1'),
        ('03', 'KM Accurio Alta Producción', 'Konica Minolta', 'AMBA', 'P2'),
        ('04', 'Insumos KM', 'Konica Minolta', 'Todo Argentina', 'P2'),
        ('05', 'UV Cama Plana (HF-3040)', 'HUENU', 'AMBA', 'P3'),
        ('05a', 'HUENU HF-6090 (LP dedicada)', 'HUENU', 'AMBA', 'P3'),
        ('06', 'UV Rollo a Rollo (otros modelos)', 'HUENU', 'AMBA', 'P3'),
        ('06a', 'HUENU HR1804UV (LP dedicada)', 'HUENU', 'AMBA', 'P3'),
        ('07', 'UV Híbrida', 'HUENU / Xenons', 'AMBA', 'P3'),
        ('08', 'DTF Textil', 'HUENU / Hanrun', 'AMBA', 'P3'),
        ('09', 'DTF UV', 'HUENU / Hanrun', 'AMBA', 'P3'),
        ('10', 'Sublimación Gran Formato', 'HUENU', 'AMBA', 'P3'),
        ('11', 'Ecosolvente + Solvente', 'HUENU', 'AMBA', 'P3'),
        ('12', 'Tenneth Mesas de Corte (otros)', 'Tenneth', 'AMBA', 'P4'),
        ('12a', 'Tenneth FC7090 (LP dedicada)', 'Tenneth', 'AMBA', 'P4'),
        ('13', 'Etiquetas — Impresión y Corte', 'Harpy/Bizpress/Duoblade/Teneth', 'AMBA', 'P4'),
        ('14', 'Insumos Gran Formato', 'Genérico', 'Todo Argentina', 'P1'),
        ('15', 'Dlican UV Cama Plana Alta Prod.', 'Dlican', 'Todo Argentina', 'P2'),
        ('16', 'Dlican UV Híbrida Alta Prod.', 'Dlican', 'Todo Argentina', 'P2'),
        ('17', 'Sinajet Mesas de Corte Alta Prod.', 'Sinajet', 'Todo Argentina', 'P2'),
        ('18', 'RISO Inkjet Alta Velocidad', 'RISO', 'Todo Argentina', 'P2'),
        ('19', 'HUENU UV Rollo a Rollo', 'HUENU', 'Argentina sin AMBA', 'P3'),
        ('20', 'HUENU UV Cama Plana', 'HUENU', 'Argentina sin AMBA', 'P3'),
        ('21', 'HUENU UV Híbrida', 'HUENU', 'Argentina sin AMBA', 'P3'),
        ('22', 'HUENU DTF Textil', 'HUENU', 'Argentina sin AMBA', 'P3'),
        ('23', 'HUENU DTF UV', 'HUENU', 'Argentina sin AMBA', 'P3'),
        ('24', 'HUENU Sublimación Textil', 'HUENU', 'Argentina sin AMBA', 'P3'),
        ('25', 'HUENU Captación Distribuidores', 'HUENU', 'Argentina sin AMBA', 'P4'),
    ],
    col_widths=[1, 5.5, 3, 3, 1.8]
)

# Detalle de cada campaña
campanas = [
    {
        'num': '01', 'nombre': 'KM Alquiler',
        'marcas': 'Konica Minolta — Bizhub B/N, Bizhub Color, AccurioPress',
        'geo': 'AMBA (CABA + GBA)',
        'objetivo': 'Captación de contratos de alquiler',
        'tipo': 'Search',
        'publico': 'Dueños de imprenta/copistería que evitan inmovilizar capital; gerentes de operaciones que buscan cuota fija con mantenimiento incluido',
        'intencion': 'Informacional-transaccional — buscan alternativa a la compra',
        'url': 'lp-km-alquiler.html',
        'keywords': [
            '[alquiler fotocopiadora konica minolta]',
            '[alquiler multifuncional para empresa]',
            '[alquiler impresora para imprenta]',
            '[fotocopiadora en alquiler CABA]',
            '[leasing impresora konica minolta]',
            '[alquiler fotocopiadora con mantenimiento]',
            '[impresora alquiler mensual empresa]',
            '+alquiler +fotocopiadora +empresas +argentina',
            '+alquiler +multifuncional +color',
            '+fotocopiadora +alquiler +mantenimiento +incluido',
            '+alquiler +impresora +produccion',
            '+renting +impresora +empresas',
            '+konica +minolta +alquiler +CABA',
            '+alquiler +equipo +impresion +empresa',
            '+plan +alquiler +konica +minolta',
            '+alquiler +bizhub +argentina',
        ],
        'negativas': 'comprar, precio de venta, segunda mano, tóner suelto, repuesto, doméstica, personal, gratis, Canon alquiler, Ricoh alquiler',
    },
    {
        'num': '02', 'nombre': 'KM Venta Pequeños/Medios',
        'marcas': 'Konica Minolta — Bizhub 558/658/958 y C558/C658',
        'geo': 'AMBA',
        'objetivo': 'Leads de venta directa — copisterías e imprentas medianas',
        'tipo': 'Search',
        'publico': 'Dueños de copistería, administradores de imprenta mediana; pain point: equipo lento, costo por copia alto, falta de soporte confiable',
        'intencion': 'Comparativa y transaccional',
        'url': 'sistemasysoluciones.com/productos/impresion/pliegos/laser-color · /laser-b-n',
        'keywords': [
            '[konica minolta bizhub precio]',
            '[fotocopiadora konica minolta]',
            '[multifuncional color A3 para imprenta]',
            '[bizhub C658 precio argentina]',
            '[bizhub 658 precio]',
            '[fotocopiadora color para imprenta precio]',
            '[impresora laser color oficina grande]',
            '+multifuncional +konica +minolta +precio',
            '+fotocopiadora +color +para +copisteria',
            '+impresora +laser +color +A3 +precio',
            '+fotocopiadora +multifuncional +empresas',
            '+konica +minolta +distribuidor +argentina',
            '+bizhub +color +precio +argentina',
            '+impresora +produccion +mediana',
            '+fotocopiadora +A3 +color +precio',
            '+multifuncional +laser +alta +velocidad',
            '+bizhub +blanco +negro +precio',
        ],
        'negativas': 'alquiler, segunda mano, tóner, repuesto, inkjet doméstica, A4 solo, gratis, home office, personal, Canon, Ricoh, Xerox',
    },
    {
        'num': '03', 'nombre': 'KM Accurio Alta Producción',
        'marcas': 'Konica Minolta — AccurioPress C3070/80, C6100, Bizhub Pro1100',
        'geo': 'AMBA',
        'objetivo': 'Leads de venta directa — imprentas industriales',
        'tipo': 'Search',
        'publico': 'Dueños e inversores de imprentas industriales; gerentes de producción de alta tirada; pain point: uptime, costo por hoja, calidad offset digital',
        'intencion': 'Investigativa-transaccional — decisión de inversión alta',
        'url': 'sistemasysoluciones.com/productos/impresion/pliegos/alta-produccion',
        'keywords': [
            '[accuriopress konica minolta]',
            '[impresora digital alta produccion]',
            '[offset digital industrial argentina]',
            '[konica minolta accurio precio]',
            '[accuriopress C6100 precio]',
            '[impresora produccion imprenta industrial]',
            '+konica +minolta +accurio +precio +argentina',
            '+impresora +produccion +100ppm',
            '+offset +digital +impresion +industrial',
            '+accuriopress +distribuidor +argentina',
            '+impresora +alta +tirada +color',
            '+impresion +digital +industrial +argentina',
            '+konica +minolta +alta +produccion',
            '+bizhub +pro +1100 +precio',
            '+impresora +digital +para +imprenta +alta',
            '+impresora +alta +velocidad +color +produccion',
        ],
        'negativas': 'alquiler, segunda mano, tóner, repuesto, doméstica, copistería pequeña, A4 personal, gratis, tutorial, bizhub pequeño',
    },
    {
        'num': '04', 'nombre': 'Insumos KM',
        'marcas': 'Konica Minolta — tóners y repuestos originales',
        'geo': 'Todo Argentina',
        'objetivo': 'Venta directa de tóners y repuestos',
        'tipo': 'Search',
        'publico': 'Encargados de compras y jefes de producción con equipos KM; pain point: desabastecimiento, repuestos no originales, demora en entrega',
        'intencion': 'Transaccional urgente — buscan proveedor local con stock',
        'url': 'sistemasysoluciones.com/productos/insumos-y-repuestos/',
        'keywords': [
            '[toner konica minolta original]',
            '[repuesto bizhub argentina]',
            '[toner accuriopress original]',
            '[drum bizhub precio]',
            '[toner konica minolta CABA]',
            '[consumibles konica minolta argentina]',
            '+toner +konica +minolta +original +precio',
            '+repuesto +bizhub +original',
            '+toner +bizhub +C658 +precio',
            '+toner +bizhub +C558 +original',
            '+insumos +konica +minolta +argentina',
            '+drum +bizhub +original',
            '+toner +konica +minolta +distribuidor',
            '+repuesto +accuriopress +original',
            '+toner +original +konica +distribuidor',
        ],
        'negativas': 'compatible, genérico, alternativo, gratis, tutorial, segunda mano, equipo completo, alquiler',
    },
    {
        'num': '05a', 'nombre': 'HUENU HF-6090 (LP dedicada)',
        'marcas': 'HUENU — HF-6090',
        'geo': 'AMBA',
        'objetivo': 'Venta directa — LP de producto específico',
        'tipo': 'Search',
        'publico': 'Dueños de taller de merchandising y personalización; pain point: imprimir sobre objetos con acabado premium',
        'intencion': 'Investigativa-transaccional',
        'url': 'lp-hf6090.html',
        'keywords': [
            '[huenu hf 6090 precio]',
            '[impresora UV cama plana 60x90]',
            '[UV flatbed merchandising precio]',
            '[impresora UV objetos rígidos pequeña]',
            '[flatbed UV 600x900 precio argentina]',
            '+huenu +hf +6090',
            '+impresora +UV +cama +plana +merchandising',
            '+flatbed +UV +vidrio +madera +acrilico',
            '+impresora +UV +directo +objetos',
            '+UV +cama +plana +90mm +altura',
        ],
        'negativas': 'Dlican industrial, segunda mano, gratis, tinta sola, alquiler, rollo',
    },
    {
        'num': '06a', 'nombre': 'HUENU HR1804UV (LP dedicada)',
        'marcas': 'HUENU — HR1804UV',
        'geo': 'AMBA',
        'objetivo': 'Venta directa — LP de producto específico',
        'tipo': 'Search',
        'publico': 'Dueños de imprenta gráfica y señalética que buscan UV rollo de producción media-alta',
        'intencion': 'Comparativa-transaccional',
        'url': 'lp-hr1804uv.html',
        'keywords': [
            '[huenu hr1804uv precio]',
            '[plotter UV rollo 1800mm precio]',
            '[UV rollo 4 cabezales epson precio]',
            '[plotter UV 1800mm con blanco y barniz]',
            '[plotter UV rollo señaletica 1800mm]',
            '+huenu +hr1804 +UV',
            '+plotter +UV +1800mm +precio',
            '+UV +rollo +4 +cabezales +epson',
            '+plotter +UV +barniz +fluorescente',
            '+impresora +UV +rollo +1800mm +argentina',
        ],
        'negativas': 'cama plana, híbrida, Dlican industrial, segunda mano, gratis, tinta sola, alquiler',
    },
    {
        'num': '12a', 'nombre': 'Tenneth FC7090 (LP dedicada)',
        'marcas': 'Tenneth — FC7090A / FC7090U',
        'geo': 'AMBA',
        'objetivo': 'Venta directa — LP de producto específico',
        'tipo': 'Search',
        'publico': 'Dueños de taller gráfico que automatizan el corte por primera vez',
        'intencion': 'Transaccional — comparan modelos de mesa de corte digital',
        'url': 'lp-teneth7090.html',
        'keywords': [
            '[tenneth fc7090 precio]',
            '[mesa de corte digital 700x900 precio]',
            '[mesa de corte CCD contorno precio]',
            '[cortadora digital stickers 700x900]',
            '[fc7090 mesa de corte argentina]',
            '+tenneth +fc7090 +precio',
            '+mesa +corte +700x900 +precio',
            '+cortadora +digital +CCD +precio',
            '+mesa +de +corte +sin +troquel +700mm',
            '+cortadora +digital +vinilo +900mm',
        ],
        'negativas': 'Sinajet, industrial, alta producción, guillotina, manual, segunda mano, gratis',
    },
    {
        'num': '14', 'nombre': 'Insumos Gran Formato',
        'marcas': 'Sin marca — compatible con cualquier equipo',
        'geo': 'Todo Argentina',
        'objetivo': 'Venta directa de tintas, films, polvos y cabezales',
        'tipo': 'Search',
        'publico': 'Operadores de plotter de cualquier marca; encargados de compras en talleres gráficos',
        'intencion': 'Transaccional — buscan proveedor, precio, disponibilidad',
        'url': 'sistemasysoluciones.com/productos/insumos-y-repuestos/tintas',
        'keywords': [
            '[tinta UV para plotter]',
            '[film DTF textil]',
            '[polvo DTF adhesivo]',
            '[cabezal Epson i3200]',
            '[tinta sublimacion para plotter]',
            '[tinta ecosolvente para plotter]',
            '[tinta UV compatible plotter]',
            '[repuesto cabezal plotter]',
            '+tinta +UV +plotter +argentina',
            '+film +DTF +rollo',
            '+cabezal +ricoh +repuesto',
            '+tinta +sublimacion +argentina',
            '+polvo +DTF +fino',
            '+insumos +plotter +gran +formato',
            '+cabezal +epson +original',
            '+tinta +ecosolvente +compatible',
            '+insumos +impresion +digital +argentina',
            '+film +DTF +mate +brillante',
        ],
        'negativas': 'gratis, tutorial, segunda mano, impresora doméstica, HP, Canon, Epson doméstica, alquiler, reparación',
    },
    {
        'num': '15', 'nombre': 'Dlican UV Cama Plana Alta Producción',
        'marcas': 'Dlican — DLI-9060 a DLI-3325',
        'geo': 'Todo Argentina',
        'objetivo': 'Venta directa — equipos industriales UV flatbed',
        'tipo': 'Search',
        'publico': 'Gerentes de producción en imprentas industriales de señalética',
        'intencion': 'Investigativa-transaccional — inversión alta',
        'url': 'sistemasysoluciones.com/productos/impresion/gran-formato/cama-plana',
        'keywords': [
            '[impresora UV cama plana industrial]',
            '[flatbed UV industrial precio argentina]',
            '[dlican UV flatbed precio]',
            '[impresora UV rigidos industrial argentina]',
            '+impresora +UV +cama +plana +industrial',
            '+flatbed +UV +industrial +argentina',
            '+plotter +UV +plano +gran +formato',
            '+dlican +flatbed +UV +precio',
            '+impresora +UV +señaletica +industrial',
            '+flatbed +UV +ricoh +precio',
        ],
        'negativas': 'HUENU pequeño, segunda mano, gratis, tinta sola, rollo, doméstica, entrada nivel',
    },
    {
        'num': '16', 'nombre': 'Dlican UV Híbrida Alta Producción',
        'marcas': 'Dlican — DLI-1688 a DLI-6600',
        'geo': 'Todo Argentina',
        'objetivo': 'Venta directa — equipos industriales UV híbridos',
        'tipo': 'Search',
        'publico': 'Imprentas industriales que producen sobre rollo y rígidos a gran escala',
        'intencion': 'Investigativa-transaccional — nicho industrial',
        'url': 'sistemasysoluciones.com/productos/impresion/gran-formato/uv-híbrido',
        'keywords': [
            '[impresora UV híbrida industrial]',
            '[dlican UV híbrido precio]',
            '[UV híbrida industrial argentina]',
            '[impresora UV rollo rigido industrial]',
            '+impresora +UV +hibrida +industrial',
            '+dlican +UV +hibrido +precio',
            '+UV +hibrida +rollo +rigido +industrial',
            '+plotter +UV +hibrido +ricoh',
            '+UV +hibrida +señaletica +industrial',
            '+dlican +hibrido +precio +argentina',
        ],
        'negativas': 'HUENU mediano, segunda mano, gratis, tinta sola, entrada nivel',
    },
    {
        'num': '17', 'nombre': 'Sinajet Mesas de Corte Alta Producción',
        'marcas': 'Sinajet — Series DG, DF-MT, DH, DF',
        'geo': 'Todo Argentina',
        'objetivo': 'Venta directa — mesas de corte industriales',
        'tipo': 'Search',
        'publico': 'Gerentes de producción en imprentas de packaging y señalética',
        'intencion': 'Transaccional — reemplazan corte manual o mesa básica por industrial',
        'url': 'sistemasysoluciones.com/productos/corte/mesas_de_corte/serie-dg · /serie-dh',
        'keywords': [
            '[mesa de corte digital industrial]',
            '[sinajet mesa de corte precio]',
            '[cortadora digital alta produccion]',
            '+mesa +de +corte +digital +industrial',
            '+sinajet +cortadora +precio',
            '+mesa +corte +señaletica +gran +formato',
            '+cortadora +digital +packaging +argentina',
            '+mesa +de +corte +para +impresion +industrial',
            '+cortadora +material +grafico +automatica',
            '+mesa +corte +alta +produccion +argentina',
        ],
        'negativas': 'manual, guillotina, Tenneth entry, segunda mano, gratis, laser, pequeño',
    },
    {
        'num': '18', 'nombre': 'RISO Inkjet Alta Velocidad',
        'marcas': 'RISO — ComColor GD7330, FW5230',
        'geo': 'Todo Argentina',
        'objetivo': 'Venta directa — impresión inkjet alta velocidad',
        'tipo': 'Search',
        'publico': 'Directores de imprentas de mailing y tiradas largas',
        'intencion': 'Comparativa-transaccional — evalúan alternativas a offset o láser',
        'url': 'sistemasysoluciones.com/productos/impresion/pliegos/inkjet',
        'keywords': [
            '[impresora riso comcolor precio]',
            '[riso impresora alta velocidad]',
            '[inkjet alta produccion argentina]',
            '[impresora 120 paginas por minuto]',
            '[impresora para tiradas largas]',
            '+impresora +riso +argentina +precio',
            '+inkjet +alta +velocidad +impresion',
            '+impresora +mailing +alta +tirada',
            '+riso +impresion +digital +precio',
            '+impresora +inkjet +produccion +industrial',
        ],
        'negativas': 'láser, tóner, segunda mano, doméstica, A4 personal, HP, Canon, Epson doméstica',
    },
    {
        'num': '19', 'nombre': 'HUENU UV Rollo a Rollo (Interior)',
        'marcas': 'HUENU — HR UV series',
        'geo': 'Argentina sin AMBA',
        'objetivo': 'Demanda usuario final + reconocimiento de marca HUENU',
        'tipo': 'Search',
        'publico': 'Dueños de estudio gráfico y señalética en el interior',
        'intencion': 'Investigativa — comparan marcas y tecnologías',
        'url': 'huenu.ar (HR1804UV publicado)',
        'keywords': [
            '[plotter UV rollo precio argentina]',
            '[impresora UV rollo señaletica]',
            '[huenu UV rollo precio]',
            '+plotter +UV +rollo +precio +argentina',
            '+impresora +UV +rollo +señaletica',
            '+plotter +UV +rollo +interior',
            '+huenu +plotter +UV',
            '+impresora +UV +rollo +1800mm',
            '+plotter +UV +vinilo +precio',
            '+impresora +UV +señaletica +interior',
        ],
        'negativas': 'segunda mano, tinta sola, alquiler, gratis, tutorial (geo AMBA excluida por segmentación)',
    },
    {
        'num': '25', 'nombre': 'HUENU Captación Distribuidores',
        'marcas': 'HUENU — toda la línea',
        'geo': 'Argentina sin AMBA',
        'objetivo': 'Sumar distribuidores a la red nacional HUENU',
        'tipo': 'Search + Display remarketing',
        'publico': 'Dueños de negocio de equipos gráficos, distribuidores de insumos',
        'intencion': 'B2B — búsqueda de oportunidades de negocio',
        'url': 'huenu.ar/Distribuidores (crear antes de activar)',
        'keywords': [
            '[distribuidor de plotters argentina]',
            '[distribuir equipos de impresión digital]',
            '[ser distribuidor de plotters]',
            '+distribuir +plotters +argentina',
            '+revendedor +equipos +impresion',
            '+distribuidor +plotter +gran +formato',
            '+sumar +linea +impresion +negocio',
            '+distribuidor +marcas +impresion',
            '+representante +equipos +graficos',
            '+distribuidor +huenu +argentina',
        ],
        'negativas': 'usuario final, precio para uso propio, gratis, tutorial',
    },
    {
        'num': '05', 'nombre': 'UV Cama Plana — HF-3040',
        'marcas': 'HUENU — HF-3040',
        'geo': 'AMBA (CABA + GBA)',
        'objetivo': 'Leads de venta — entrada al UV sobre objetos rígidos',
        'tipo': 'Search',
        'publico': 'Emprendedores y talleres pequeños que quieren empezar en impresión UV; pain point: presupuesto limitado, buscan el modelo de entrada',
        'intencion': 'Transaccional — buscan el equipo más accesible para comenzar en UV',
        'url': 'sistemasysoluciones.com/productos/impresion/gran-formato/cama-plana',
        'keywords': [
            '[impresora UV cama plana precio]',
            '[plotter UV flatbed entrada nivel]',
            '[impresora UV objetos pequeña]',
            '[HUENU HF-3040]',
            '[impresora UV A3 precio argentina]',
            '+impresora +UV +cama +plana +precio',
            '+plotter +UV +para +madera +vidrio',
            '+impresora +UV +rigidos +entrada',
            '+UV +flatbed +pequeño +taller',
            '+impresion +UV +sobre +objetos +precio',
            '+HUENU +UV +cama +plana',
            '+plotter +UV +merchandising +precio',
            '+reseller +HUENU +argentina',
            '+impresora +UV +led +compacta',
            '+UV +flatbed +para +iniciar',
        ],
        'negativas': 'segunda mano, gratis, tinta sola, alquiler, rollo, doméstica, HF-6090',
    },
    {
        'num': '06', 'nombre': 'UV Rollo a Rollo — Otros modelos',
        'marcas': 'HUENU — HR1602UV, HR2008UV, HR3204UV, HR3208UV, HR3212UV',
        'geo': 'AMBA (CABA + GBA)',
        'objetivo': 'Leads de venta — señalética UV a gran escala',
        'tipo': 'Search',
        'publico': 'Dueños de imprenta de señalética y comunicación visual; gerentes de producción; pain point: calidad UV en grandes tiradas, ancho de trabajo',
        'intencion': 'Comparativa y transaccional — buscan mayor capacidad que HR1804UV',
        'url': 'sistemasysoluciones.com/productos/impresion/gran-formato/uv-rollo-a-rollo',
        'keywords': [
            '[plotter UV rollo a rollo precio]',
            '[plotter UV gran formato señaletica]',
            '[impresora UV 2000mm precio]',
            '[plotter UV 3200mm argentina]',
            '[HUENU UV rollo a rollo]',
            '+plotter +UV +rollo +gran +formato',
            '+impresora +UV +señaletica +industrial',
            '+UV +rollo +a +rollo +precio +argentina',
            '+plotter +UV +alta +produccion',
            '+impresion +UV +exterior +gran +formato',
            '+HUENU +UV +rollo +precio',
            '+plotter +UV +3200mm',
            '+reseller +HUENU +UV',
            '+impresora +UV +para +imprenta',
            '+plotter +UV +blanco +barniz +rollo',
        ],
        'negativas': 'cama plana, segunda mano, gratis, tinta sola, alquiler, doméstica, HR1804UV',
    },
    {
        'num': '07', 'nombre': 'UV Híbrida AMBA',
        'marcas': 'HUENU HI1804UV / Xenons R180 PRO',
        'geo': 'AMBA (CABA + GBA)',
        'objetivo': 'Leads de venta — plotter UV que imprime en rollo y rígidos',
        'tipo': 'Search',
        'publico': 'Imprentas que necesitan versatilidad — señalética en vinilo y sobre rígidos sin cambiar de equipo; pain point: costo de tener dos máquinas',
        'intencion': 'Comparativa — buscan solución todo en uno para UV',
        'url': 'sistemasysoluciones.com/productos/impresion/gran-formato/uv-hibrido',
        'keywords': [
            '[impresora UV híbrida precio]',
            '[plotter UV rollo y rigidos]',
            '[UV híbrida rollo cama plana]',
            '[HUENU HI1804UV precio]',
            '[Xenons R180 PRO precio argentina]',
            '+impresora +UV +hibrida +precio',
            '+plotter +UV +rollo +y +rigidos',
            '+UV +hibrida +para +señaletica',
            '+impresora +UV +versatil +precio',
            '+HUENU +UV +hibrida +argentina',
            '+plotter +UV +dos +formatos',
            '+UV +rigidos +y +vinilo',
            '+impresora +UV +packaging +señaletica',
            '+reseller +UV +hibrida +argentina',
            '+plotter +hibrido +UV +led',
        ],
        'negativas': 'cama plana solo, rollo solo, segunda mano, gratis, tinta sola, alquiler, doméstica',
    },
    {
        'num': '08', 'nombre': 'DTF Textil AMBA',
        'marcas': 'HUENU HR-DT series / Hanrun Super A-602',
        'geo': 'AMBA (CABA + GBA)',
        'objetivo': 'Leads de venta — plotter DTF para estampado textil',
        'tipo': 'Search',
        'publico': 'Dueños de taller textil y estampería; emprendedores en remeras; pain point: calidad de DTF, shaker integrado, volumen de producción',
        'intencion': 'Transaccional — buscan equipo DTF con shaker integrado',
        'url': 'sistemasysoluciones.com/productos/impresion/dtf/dtf-textil',
        'keywords': [
            '[plotter DTF textil precio argentina]',
            '[impresora DTF para remeras]',
            '[máquina DTF textil precio]',
            '[HUENU DTF textil]',
            '[DTF con shaker integrado]',
            '+plotter +DTF +textil +precio',
            '+impresora +DTF +remeras +argentina',
            '+DTF +textil +shaker +integrado',
            '+maquina +DTF +estampado +textil',
            '+plotter +DTF +alta +produccion',
            '+DTF +textil +i3200 +precio',
            '+HUENU +DTF +textil +argentina',
            '+impresora +estampado +textil +digital',
            '+DTF +remeras +precio +argentina',
            '+plotter +DTF +profesional',
        ],
        'negativas': 'sublimación, UV, ecosolvente, segunda mano, gratis, tinta sola, alquiler, doméstica, serigrafía, bordado',
    },
    {
        'num': '09', 'nombre': 'DTF UV AMBA',
        'marcas': 'HUENU HR-DU series / Hanrun UV',
        'geo': 'AMBA (CABA + GBA)',
        'objetivo': 'Leads de venta — DTF UV para objetos y packaging',
        'tipo': 'Search',
        'publico': 'Talleres de personalización de objetos y packaging; pain point: acabado premium, adherencia sobre rígidos y flexibles sin impresora UV flatbed',
        'intencion': 'Transaccional — buscan alternativa al UV flatbed para objetos pequeños',
        'url': 'sistemasysoluciones.com/productos/impresion/dtf/dtf-uv',
        'keywords': [
            '[impresora DTF UV precio]',
            '[DTF UV para objetos precio argentina]',
            '[plotter DTF UV con laminado]',
            '[HUENU DTF UV]',
            '[DTF UV sobre rigidos precio]',
            '+impresora +DTF +UV +precio',
            '+DTF +UV +laminado +integrado',
            '+DTF +UV +packaging +objetos',
            '+plotter +DTF +UV +argentina',
            '+HUENU +DTF +UV +precio',
            '+impresora +DTF +UV +i3200',
            '+DTF +UV +merchandising +precio',
            '+plotter +DTF +UV +rollo',
            '+DTF +UV +alta +calidad',
            '+impresora +transfer +UV +precio',
        ],
        'negativas': 'DTF textil, sublimación, ecosolvente, segunda mano, gratis, tinta sola, alquiler, doméstica',
    },
    {
        'num': '10', 'nombre': 'Sublimación Gran Formato AMBA',
        'marcas': 'HUENU HR-SU series',
        'geo': 'AMBA (CABA + GBA)',
        'objetivo': 'Leads de venta — plotter de sublimación textil gran formato',
        'tipo': 'Search',
        'publico': 'Talleres de sublimación textil; productores de ropa deportiva y telas; pain point: velocidad, uniformidad de color, ancho de trabajo',
        'intencion': 'Comparativa y transaccional',
        'url': 'sistemasysoluciones.com/productos/impresion/gran-formato/sublimacion',
        'keywords': [
            '[plotter sublimacion precio argentina]',
            '[impresora sublimacion textil gran formato]',
            '[plotter sublimacion 1800mm precio]',
            '[HUENU sublimacion]',
            '[maquina sublimacion tela precio]',
            '+plotter +sublimacion +textil +precio',
            '+impresora +sublimacion +gran +formato',
            '+sublimacion +textil +industrial +precio',
            '+plotter +sublimacion +i3200',
            '+HUENU +sublimacion +precio +argentina',
            '+impresora +sublimacion +alta +produccion',
            '+plotter +sublimacion +para +tela',
            '+reseller +sublimacion +argentina',
            '+sublimacion +1800mm +precio',
            '+plotter +sublimacion +secado +ir',
        ],
        'negativas': 'DTF, UV, ecosolvente, segunda mano, gratis, tinta sola, alquiler, taza pequeña, doméstica, inkjet A4',
    },
    {
        'num': '11', 'nombre': 'Ecosolvente + Solvente Gran Formato',
        'marcas': 'HUENU HR-ES / HR-ASV series',
        'geo': 'AMBA (CABA + GBA)',
        'objetivo': 'Leads de venta — plotter ecosolvente para señalética exterior',
        'tipo': 'Search',
        'publico': 'Imprentas de señalética y banners; talleres de vinilos para exteriores; pain point: durabilidad exterior, tintas con resistencia UV, ancho de trabajo',
        'intencion': 'Transaccional — buscan plotter para banners y vinilos durables',
        'url': 'sistemasysoluciones.com/productos/impresion/gran-formato/ecosolvente',
        'keywords': [
            '[plotter ecosolvente precio argentina]',
            '[impresora ecosolvente señaletica]',
            '[plotter banners vinilos precio]',
            '[HUENU ecosolvente precio]',
            '[plotter ecosolvente 1800mm precio]',
            '+plotter +ecosolvente +precio +argentina',
            '+impresora +ecosolvente +señaletica',
            '+plotter +vinilo +exterior +precio',
            '+ecosolvente +gran +formato +argentina',
            '+HUENU +ecosolvente +precio',
            '+plotter +banners +durables',
            '+impresora +ecosolvente +i3200',
            '+plotter +ecosolvente +3200mm',
            '+señaletica +exterior +plotter',
            '+reseller +ecosolvente +argentina',
        ],
        'negativas': 'UV, sublimación, DTF, segunda mano, gratis, tinta sola, alquiler, doméstica, inkjet A4',
    },
    {
        'num': '12', 'nombre': 'Tenneth — Otros Modelos',
        'marcas': 'Tenneth FC9012, FC1313U',
        'geo': 'AMBA (CABA + GBA)',
        'objetivo': 'Leads de venta — mesas de corte digital para mayor volumen',
        'tipo': 'Search',
        'publico': 'Imprentas con mayor volumen de corte que buscan modelos más grandes; pain point: velocidad, cama de trabajo, automatización con conveyor',
        'intencion': 'Comparativa — ya conocen Tenneth, buscan capacidad superior a FC7090',
        'url': 'sistemasysoluciones.com/productos/corte/mesas_de_corte/equipos-compatos-corte',
        'keywords': [
            '[mesa de corte digital Tenneth]',
            '[Tenneth FC9012 precio]',
            '[cortadora digital gran formato]',
            '[mesa de corte con conveyor precio]',
            '[Tenneth FC1313U precio argentina]',
            '+mesa +de +corte +digital +precio',
            '+cortadora +digital +con +transportadora',
            '+mesa +corte +1300mm +precio',
            '+Tenneth +mesa +corte +argentina',
            '+distribuidor +Tenneth +argentina',
            '+mesa +corte +digital +alta +produccion',
            '+cortadora +CNC +impresion +grafica',
            '+mesa +corte +cámara +CCD',
            '+corte +digital +packaging',
            '+Tenneth +precio +argentina',
        ],
        'negativas': 'Sinajet, manual, guillotina, FC7090, segunda mano, gratis, laser, pequeño',
    },
    {
        'num': '13', 'nombre': 'Etiquetas — Impresión y Corte',
        'marcas': 'Harpy, Bizpress, Duoblade, Teneth RN3',
        'geo': 'AMBA (CABA + GBA)',
        'objetivo': 'Leads de venta — sistemas de impresión y corte de etiquetas in-house',
        'tipo': 'Search',
        'publico': 'Empresas de packaging, logística y alimentos que tercerizan etiquetas; pain point: dependencia de proveedor, tiradas cortas, cambios de diseño frecuentes',
        'intencion': 'Investigativa-transaccional — evaluando producción in-house de etiquetas',
        'url': 'sistemasysoluciones.com/productos/impresion/etiquetas',
        'keywords': [
            '[impresora de etiquetas autoadhesivas]',
            '[máquina para hacer etiquetas precio]',
            '[cortadora de etiquetas digital precio]',
            '[Harpy impresora etiquetas argentina]',
            '[Duoblade troqueladora etiquetas]',
            '+impresora +etiquetas +rollo +precio',
            '+cortadora +etiquetas +digital',
            '+produccion +etiquetas +in-house',
            '+sistema +etiquetas +autoadhesivas',
            '+impresora +inkjet +etiquetas +rollo',
            '+Harpy +R440 +precio',
            '+Bizpress +etiquetas +laser',
            '+etiquetas +packaging +produccion +propia',
            '+troqueladora +etiquetas +digital',
            '+impresion +corte +etiquetas +precio',
        ],
        'negativas': 'gran formato, banner, lona, segunda mano, gratis, doméstica, A4, manual, etiqueta de tela',
    },
    {
        'num': '20', 'nombre': 'HUENU UV Cama Plana — Interior',
        'marcas': 'HUENU — HF-3040, HF-6090',
        'geo': 'Argentina sin AMBA',
        'objetivo': 'Captación de leads en interior para UV cama plana',
        'tipo': 'Search',
        'publico': 'Talleres y emprendedores del interior que buscan equipo UV; pain point: falta de distribuidor local, soporte a distancia',
        'intencion': 'Transaccional — buscan dónde conseguir UV cama plana fuera de AMBA',
        'url': 'sistemasysoluciones.com/productos/impresion/gran-formato/cama-plana',
        'keywords': [
            '[impresora UV cama plana argentina]',
            '[plotter UV flatbed interior argentina]',
            '[HUENU UV cama plana precio]',
            '[distribuidor HUENU interior]',
            '[impresora UV sobre objetos argentina]',
            '+HUENU +UV +cama +plana +interior',
            '+plotter +UV +flatbed +argentina',
            '+impresora +UV +rigidos +interior',
            '+distribuidor +HUENU +interior +argentina',
            '+UV +cama +plana +donde +comprar',
            '+HUENU +HF-6090 +precio',
            '+impresora +UV +objetos +argentina',
            '+plotter +UV +merchandising +interior',
            '+HUENU +distribuidor +interior',
            '+UV +flatbed +argentina +precio',
        ],
        'negativas': 'segunda mano, gratis, tinta sola, alquiler, rollo, doméstica (geo AMBA excluida por segmentación)',
    },
    {
        'num': '21', 'nombre': 'HUENU UV Híbrida — Interior',
        'marcas': 'HUENU HI1804UV',
        'geo': 'Argentina sin AMBA',
        'objetivo': 'Captación de leads en interior para UV híbrida',
        'tipo': 'Search',
        'publico': 'Imprentas del interior que necesitan versatilidad UV; pain point: un solo equipo para rollo y rígidos',
        'intencion': 'Comparativa — buscan solución todo en uno fuera de AMBA',
        'url': 'sistemasysoluciones.com/productos/impresion/gran-formato/uv-hibrido',
        'keywords': [
            '[impresora UV híbrida argentina]',
            '[plotter UV rollo y rigidos interior]',
            '[HUENU HI1804UV precio]',
            '[UV híbrida distribuidor interior]',
            '[plotter UV híbrido argentina precio]',
            '+HUENU +UV +hibrida +interior',
            '+plotter +UV +hibrida +argentina',
            '+UV +rollo +rigidos +interior',
            '+HUENU +hibrida +distribuidor',
            '+impresora +UV +versatil +interior',
            '+UV +hibrida +señaletica +interior',
            '+plotter +UV +dos +formatos +argentina',
            '+HUENU +HI1804UV',
            '+UV +híbrida +donde +comprar',
            '+reseller +HUENU +UV +interior',
        ],
        'negativas': 'cama plana solo, rollo solo, segunda mano, gratis, tinta sola, alquiler (geo AMBA excluida por segmentación)',
    },
    {
        'num': '22', 'nombre': 'HUENU DTF Textil — Interior',
        'marcas': 'HUENU HR-DT series',
        'geo': 'Argentina sin AMBA',
        'objetivo': 'Captación de leads en interior para DTF textil',
        'tipo': 'Search',
        'publico': 'Talleres textiles del interior; pain point: conseguir equipo DTF con soporte fuera de AMBA',
        'intencion': 'Transaccional — buscan DTF con respaldo local en el interior',
        'url': 'sistemasysoluciones.com/productos/impresion/dtf/dtf-textil',
        'keywords': [
            '[plotter DTF textil interior argentina]',
            '[DTF remeras interior argentina]',
            '[HUENU DTF textil distribuidor]',
            '[máquina DTF interior precio]',
            '[DTF textil donde comprar argentina]',
            '+HUENU +DTF +textil +interior',
            '+plotter +DTF +interior +argentina',
            '+DTF +remeras +interior',
            '+distribuidor +HUENU +DTF +interior',
            '+DTF +shaker +integrado +interior',
            '+plotter +DTF +provincia +argentina',
            '+HUENU +DTF +precio +interior',
            '+estampado +DTF +interior',
            '+DTF +textil +donde +comprar',
            '+plotter +DTF +argentina +precio',
        ],
        'negativas': 'sublimación, UV, ecosolvente, segunda mano, gratis, tinta sola, alquiler (geo AMBA excluida por segmentación)',
    },
    {
        'num': '23', 'nombre': 'HUENU DTF UV — Interior',
        'marcas': 'HUENU HR-DU series',
        'geo': 'Argentina sin AMBA',
        'objetivo': 'Captación de leads en interior para DTF UV',
        'tipo': 'Search',
        'publico': 'Talleres de personalización del interior; pain point: DTF UV con soporte local fuera de AMBA',
        'intencion': 'Transaccional — buscan DTF UV con respaldo en el interior',
        'url': 'sistemasysoluciones.com/productos/impresion/dtf/dtf-uv',
        'keywords': [
            '[impresora DTF UV interior argentina]',
            '[DTF UV objetos interior precio]',
            '[HUENU DTF UV distribuidor]',
            '[DTF UV donde comprar argentina]',
            '[plotter DTF UV interior]',
            '+HUENU +DTF +UV +interior',
            '+DTF +UV +laminado +interior',
            '+DTF +UV +interior +argentina',
            '+distribuidor +HUENU +DTF +UV',
            '+plotter +DTF +UV +precio +interior',
            '+DTF +UV +objetos +interior',
            '+HUENU +HR6004DU +interior',
            '+DTF +UV +donde +comprar',
            '+DTF +UV +packaging +interior',
            '+plotter +DTF +UV +argentina',
        ],
        'negativas': 'DTF textil, sublimación, ecosolvente, segunda mano, gratis, alquiler (geo AMBA excluida por segmentación)',
    },
    {
        'num': '24', 'nombre': 'HUENU Sublimación Textil — Interior',
        'marcas': 'HUENU HR-SU series',
        'geo': 'Argentina sin AMBA',
        'objetivo': 'Captación de leads en interior para sublimación textil',
        'tipo': 'Search',
        'publico': 'Talleres textiles del interior que subliman telas; pain point: soporte técnico local, repuestos disponibles fuera de AMBA',
        'intencion': 'Transaccional — buscan plotter sublimación con respaldo en el interior',
        'url': 'sistemasysoluciones.com/productos/impresion/gran-formato/sublimacion',
        'keywords': [
            '[plotter sublimacion interior argentina]',
            '[sublimacion textil interior argentina]',
            '[HUENU sublimacion distribuidor interior]',
            '[maquina sublimacion donde comprar argentina]',
            '[plotter sublimacion interior precio]',
            '+HUENU +sublimacion +interior',
            '+plotter +sublimacion +interior +argentina',
            '+sublimacion +textil +interior',
            '+distribuidor +HUENU +sublimacion',
            '+sublimacion +gran +formato +interior',
            '+plotter +sublimacion +donde +comprar',
            '+HUENU +sublimacion +precio +interior',
            '+sublimacion +tela +interior',
            '+plotter +sublimacion +provincia',
            '+sublimacion +textil +argentina',
        ],
        'negativas': 'DTF, UV, ecosolvente, segunda mano, gratis, tinta sola, alquiler (geo AMBA excluida por segmentación)',
    },
]

page_break(doc)
add_heading(doc, 'Detalle de campañas principales', 2)
add_body(doc, 'Se detallan a continuación las campañas más relevantes. El resto sigue el mismo esquema de campos.', italic=True)

for c in campanas:
    doc.add_paragraph()
    add_heading(doc, f"Campaña {c['num']} · {c['nombre']}", 3)
    make_table(doc,
        ['Campo', 'Detalle'],
        [
            ('Marcas', c['marcas']),
            ('Geografía', c['geo']),
            ('Objetivo', c['objetivo']),
            ('Tipo', c['tipo']),
            ('Público objetivo', c['publico']),
            ('Intención de búsqueda', c['intencion']),
            ('URL de destino', c['url']),
            ('Keywords negativas', c['negativas']),
        ],
        col_widths=[4, 12]
    )
    add_heading(doc, 'Keywords principales', 4)
    for kw in c['keywords']:
        add_bullet(doc, kw, size=9)

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# FASE 3 — AVISOS RSA (resumen por campaña)
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, 'FASE 3 — AVISOS RSA Y EXTENSIONES', 1)
add_body(doc, 'Formato: Anuncio de Búsqueda Responsivo (RSA) · 15 títulos (≤30 car.) · 4 descripciones (≤90 car.)', bold=True)
add_body(doc, 'Todos los avisos incluyen: 1 título con keyword principal · 1 título con diferencial (40 años / soporte / oficial) · 1 título con CTA.')
add_body(doc, 'Campañas HUENU interior: incluyen "Pedísela a tu distribuidor" como CTA secundario en al menos una descripción.')
doc.add_paragraph()

avisos = [
    {
        'num': '01', 'nombre': 'KM Alquiler',
        'titulos': [
            'Alquiler Konica Minolta', 'Sin inversión inicial', 'Cuota fija mensual',
            'Mantenimiento incluido', 'BiZhub y AccurioPress', 'Soporte técnico local',
            '40 años en Argentina', 'Consultá planes hoy', 'Imprimí sin comprar',
            'CABA y Gran Buenos Aires', 'Propuesta a medida', 'Konica Minolta Oficial',
            'Pedí tu cotización ya', 'Escalá tu producción', 'Sin gastos de reparación',
        ],
        'descripciones': [
            'Cuota fija sin inversión inicial. Mantenimiento incluido. Consultá planes disponibles.',
            'BiZhub y AccurioPress disponibles. CABA y GBA. Sin costos de reparación inesperados.',
            'Más de 40 años en el mercado argentino. Técnicos especializados. Respondemos en 24hs.',
            'Analizamos tu volumen y te recomendamos el equipo ideal. Propuesta sin compromiso.',
        ],
        'url': 'sistemasysoluciones.com/KM-Alquiler',
        'sitelinks': [
            ('Ver modelos disponibles', 'BiZhub B/N y Color', 'AccurioPress alta producción'),
            ('Cómo funciona el alquiler', 'Consulta, propuesta, instalación', 'Soporte durante el contrato'),
            ('Pedí tu propuesta', 'Sin compromiso ni costo', 'Asesor disponible en 24hs'),
            ('Contacto directo', 'Tel. y WhatsApp disponibles', 'Belgrano 748 — CABA'),
        ],
        'callouts': ['40 años en el mercado', 'Soporte técnico propio', 'Sin inversión inicial', 'Respuesta en 24 horas'],
        'snippet': ('Modelos', 'BiZhub 558, BiZhub 658, BiZhub C558, BiZhub C658, AccurioPress C3070'),
    },
    {
        'num': '02', 'nombre': 'KM Venta Pequeños/Medios',
        'titulos': [
            'Konica Minolta Oficial', 'Multifuncional Color A3', 'Fotocopiadora para imprenta',
            'BiZhub Color y B/N', 'Soporte técnico local', '40 años en Argentina',
            'Pedí tu cotización', 'Financiación disponible', 'Bizhub C658 y C558',
            'Alta velocidad de impresión', 'Para copisterías e imprentas', 'Stock disponible en CABA',
            'Consultá precio y modelos', 'Impresión laser color A3', 'Asesoramiento sin cargo',
        ],
        'descripciones': [
            'Multifuncionales Konica Minolta color y B/N. Distribuidores oficiales en CABA y GBA.',
            'BiZhub C558, C658, 658 y 958. Alta velocidad y calidad profesional para tu negocio.',
            '40 años en el mercado argentino. Soporte técnico propio. Repuestos disponibles.',
            'Consultá precio, financiación y disponibilidad. Te respondemos en menos de 24hs.',
        ],
        'url': 'sistemasysoluciones.com/Konica-Minolta',
        'sitelinks': [
            ('Ver modelos color', 'Bizhub C558 y C658', 'Hasta 65 ppm en color'),
            ('Ver modelos B/N', 'Bizhub 558, 658 y 958', 'Hasta 95 ppm en B/N'),
            ('Soporte postventa', 'Técnicos especializados', 'Atención en CABA y GBA'),
            ('Contacto y presupuesto', 'Sin compromiso ni costo', 'Belgrano 748 — CABA'),
        ],
        'callouts': ['40 años en el mercado', 'Distribuidor oficial KM', 'Soporte técnico propio', 'Repuestos en stock'],
        'snippet': ('Modelos', 'Bizhub 558, Bizhub 658, Bizhub 958, Bizhub C558, Bizhub C658'),
    },
    {
        'num': '05a', 'nombre': 'HUENU HF-6090',
        'titulos': [
            'Impresora UV Cama Plana', 'HUENU HF-6090 Oficial', 'Imprimí sobre objetos rígidos',
            'Madera, vidrio y acrílico', 'UV LED con blanco y barniz', 'Soporte técnico local',
            '40 años en Argentina', 'Cotizá tu equipo hoy', 'Para merchandising y deco',
            'Impresora UV 60×90 cm', 'Reseller oficial HUENU', 'Capacitación incluida',
            'Repuestos en el país', 'Consultá precio y stock', 'Alta definición 1800 dpi',
        ],
        'descripciones': [
            'HUENU HF-6090: impresión UV directa sobre madera, vidrio, acrílico y más. 600×900mm.',
            'CMYK + Blanco + Barniz. Hasta 1800 dpi. Medición automática de altura. Soporte local.',
            'Reseller oficial HUENU en Argentina. Capacitación y repuestos incluidos. 40 años.',
            'Ideal para merchandising, packaging y personalización. Consultá precio y financiación.',
        ],
        'url': 'sistemasysoluciones.com/UV-Cama-Plana',
        'sitelinks': [
            ('HUENU HF-6090', '600×900mm, 1800 dpi', 'CMYK+Blanco+Barniz+UV LED'),
            ('HUENU HF-3040', '297×420mm compacto', 'Ideal para comenzar en UV'),
            ('Materiales compatibles', 'Madera, vidrio, acrílico', 'Metal, cuero, packaging'),
            ('Pedí tu cotización', 'Sin compromiso ni costo', 'Respondemos en 24hs'),
        ],
        'callouts': ['Reseller oficial HUENU', 'Soporte técnico local', 'Capacitación incluida', 'Repuestos en stock'],
        'snippet': ('Aplicaciones', 'Merchandising, Packaging, Señalética, Arte, Vidrio, Madera'),
    },
    {
        'num': '06a', 'nombre': 'HUENU HR1804UV',
        'titulos': [
            'Plotter UV Rollo a Rollo', 'HUENU HR1804 UV Oficial', 'Señalética y comunicación',
            'CMYK, Blanco y Barniz', '1800mm de ancho', '4 cabezales Epson I3200',
            'Soporte técnico local', '40 años en Argentina', 'Cotizá tu plotter UV',
            'Para vinilo y wallpapers', 'Reseller oficial HUENU', 'Capacitación incluida',
            'Alta producción UV', 'Consultá modelos y precio', 'Repuestos en el país',
        ],
        'descripciones': [
            'HUENU HR1804UV: 1800mm, 4 cabezales Epson I3200, CMYK+Blanco+Barniz+Flúo.',
            'Impresión UV en vinilo, wallpapers, gráficos vehiculares y comunicación visual.',
            'Reseller oficial HUENU en Argentina. Capacitación, soporte técnico y repuestos.',
            '40 años en el mercado argentino. Consultá el modelo ideal para tu volumen.',
        ],
        'url': 'sistemasysoluciones.com/UV-Rollo',
        'sitelinks': [
            ('HUENU HR1804 UV', '1800mm, 4 cabezales Epson', 'CMYK+Blanco+Barniz+Flúo'),
            ('Ver toda la línea UV', 'Desde 1600mm hasta 3200mm', 'Hasta 12 cabezales'),
            ('Aplicaciones UV', 'Vinilo, wallpaper, wrapping', 'Comunicación visual'),
            ('Pedí tu cotización', 'Sin compromiso ni costo', 'Respondemos en 24hs'),
        ],
        'callouts': ['Reseller oficial HUENU', 'Soporte técnico local', 'Capacitación incluida', 'Repuestos en stock'],
        'snippet': ('Modelos', 'HR1602UV, HR1804UV, HR2008UV, HR3204UV, HR3208UV, HR3212UV'),
    },
    {
        'num': '12a', 'nombre': 'Tenneth FC7090',
        'titulos': [
            'Mesa de Corte Digital', 'Tenneth FC7090 Oficial', 'Corte de contorno CCD',
            'Precisión ±0,1mm', 'Sin troquel, sin desperdicio', 'Soporte técnico local',
            '40 años en Argentina', 'Cotizá tu mesa de corte', 'Stickers y etiquetas',
            'Hasta 1000mm/s', 'Distribuidor oficial', 'Capacitación incluida',
            'Consultá modelos y precio', 'Corte y hendido en uno', 'Para vinilo y POP',
        ],
        'descripciones': [
            'Tenneth FC7090: 700×900mm, cámara CCD, precisión ±0,1mm. Hasta 1000mm/s.',
            'Corte de stickers, etiquetas, vinilo y POP sin troquel. Doble herramienta.',
            'Distribuidor oficial Tenneth en Argentina. Capacitación incluida. Soporte local.',
            '40 años en el mercado argentino. Automatizá tu corte y escalá sin más personal.',
        ],
        'url': 'sistemasysoluciones.com/Mesa-Corte',
        'sitelinks': [
            ('Tenneth FC7090', '700×900mm, CCD contorno', 'Mesa fija o transportadora'),
            ('Tenneth FC1313U', '1300×1300mm conveyor', 'Para mayor volumen de corte'),
            ('Tenneth FC5070E', '500×700mm compacto', 'Entrada al corte digital'),
            ('Pedí tu cotización', 'Sin compromiso ni costo', 'Respondemos en 24hs'),
        ],
        'callouts': ['Cámara CCD de contorno', 'Precisión ±0,1mm', 'Soporte técnico local', 'Capacitación incluida'],
        'snippet': ('Modelos', 'FC5070E, FC7090A, FC7090U, FC9012, FC1313U'),
    },
    {
        'num': '19', 'nombre': 'HUENU UV Rollo (Interior)',
        'titulos': [
            'Plotter UV HUENU', 'UV Rollo a Rollo HUENU', 'HUENU HR1804UV',
            'Señalética con calidad UV', 'CMYK, Blanco y Barniz', '4 cabezales Epson I3200',
            'Soporte local en Argentina', 'Pedísela a tu distribuidor', 'Para vinilo y wallpapers',
            '1800mm de ancho', 'Marca HUENU en Argentina', 'Red de distribuidores',
            'Consultá donde conseguirla', 'Impresión UV profesional', 'Calidad con respaldo local',
        ],
        'descripciones': [
            'HUENU HR1804UV: 1800mm, 4 cabezales Epson I3200, CMYK+Blanco+Barniz. Alta calidad.',
            'Señalética, viniles, wallpapers y comunicación visual con tecnología UV profesional.',
            'HUENU tiene distribuidores en todo el interior del país. Pedísela a tu proveedor local.',
            'Soporte técnico y repuestos disponibles en Argentina. Consultá dónde conseguir tu HUENU.',
        ],
        'url': 'huenu.ar/UV-Rollo',
        'sitelinks': [
            ('HUENU HR1804UV', '1800mm, Epson I3200', 'CMYK+Blanco+Barniz+Flúo'),
            ('Ver toda la línea UV', 'Desde 1600 a 3200mm', 'Hasta 12 cabezales'),
            ('Red de distribuidores', 'Encontrá tu distribuidor', 'En todo el interior'),
            ('Consultá tu modelo', 'Sin compromiso ni costo', 'Respondemos en 24hs'),
        ],
        'callouts': ['Marca con respaldo local', 'Red de distribuidores', 'Repuestos en Argentina', 'Soporte técnico'],
        'snippet': ('Modelos', 'HR1602UV, HR1804UV, HR2008UV, HR3204UV, HR3208UV, HR3212UV'),
    },
    {
        'num': '25', 'nombre': 'HUENU Captación Distribuidores',
        'titulos': [
            'Distribuí HUENU en tu zona', 'Sumá una línea rentable', 'Representá la marca HUENU',
            'Red de distribuidores', 'Margen y respaldo de marca', 'Soporte comercial incluido',
            'HUENU — marca en crecimiento', 'Oportunidad de negocio', 'Para el interior del país',
            'Equipos de alta demanda', 'Sumate a la red HUENU', 'Capacitación y materiales',
            'Consultá condiciones hoy', 'Negocio con respaldo real', 'Pedí info sin compromiso',
        ],
        'descripciones': [
            'Distribuí equipos HUENU en tu zona. Línea completa: UV, DTF, sublimación y ecosolvente.',
            'Soporte comercial, capacitación y materiales de marca incluidos. Margen competitivo.',
            'Marca con respaldo técnico local. Repuestos en el país. Red nacional en crecimiento.',
            'Consultá condiciones para distribuidores en el interior de Argentina. Sin compromiso.',
        ],
        'url': 'huenu.ar/Distribuidores',
        'sitelinks': [
            ('Línea de productos', 'UV, DTF, sublimación', 'Ecosolvente y corte'),
            ('Soporte al distribuidor', 'Capacitación incluida', 'Materiales de marca'),
            ('Condiciones comerciales', 'Margen competitivo', 'Respaldo técnico local'),
            ('Consultá disponibilidad', 'Por zona, sin compromiso', 'Respondemos en 24hs'),
        ],
        'callouts': ['Margen competitivo', 'Soporte técnico incluido', 'Capacitación y materiales', 'Red en crecimiento'],
        'snippet': ('Líneas', 'UV Rollo, UV Cama Plana, DTF Textil, DTF UV, Sublimación, Ecosolvente'),
    },
]

for av in avisos:
    doc.add_paragraph()
    add_heading(doc, f"Campaña {av['num']} · {av['nombre']}", 3)
    add_heading(doc, 'RSA — 15 Títulos (≤30 caracteres)', 4)
    for i, t in enumerate(av['titulos'], 1):
        add_bullet(doc, f"{i:02d}. {t}", size=9)
    add_heading(doc, '4 Descripciones (≤90 caracteres)', 4)
    for i, d in enumerate(av['descripciones'], 1):
        add_bullet(doc, f"D{i}: {d}", size=9)
    add_body(doc, f"URL visible: {av['url']}", bold=True, size=9)
    add_heading(doc, 'Extensiones de sitio', 4)
    make_table(doc,
        ['Texto', 'Descripción 1', 'Descripción 2'],
        av['sitelinks'],
        col_widths=[4, 5, 5]
    )
    add_body(doc, 'Llamada: (011) 4342-5742', size=9)
    add_body(doc, 'Textos destacados: ' + ' · '.join(av['callouts']), size=9)
    add_body(doc, f"Snippet — {av['snippet'][0]}: {av['snippet'][1]}", size=9)

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# FASE 4 — PUJAS Y PRESUPUESTO
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, 'FASE 4 — PUJAS Y PRESUPUESTO', 1)
add_body(doc, 'Tipo de cambio de referencia: USD 1 = ARS 1.200 — verificar y actualizar al momento de activación.', bold=True)
doc.add_paragraph()

add_heading(doc, 'Estrategia de pujas', 2)
add_body(doc, 'Fase inicial: Maximizar clics con CPC máximo por campaña.')
add_body(doc, 'Objetivo: acumular historial de conversiones antes de activar smart bidding.')
doc.add_paragraph()
make_table(doc,
    ['Condición', 'Acción'],
    [
        ('≥ 30 conversiones/mes por campaña', 'Migrar a CPA objetivo'),
        ('< 15 conversiones en 60 días', 'Revisar keywords, landing y presupuesto antes de migrar'),
        ('KM Alquiler con > 10 contratos/mes', 'Evaluar Maximizar conversiones sin techo'),
        ('Campañas HUENU interior con bajo volumen', 'Mantener Maximizar clics hasta acumular historial'),
    ],
    col_widths=[8, 8]
)

add_heading(doc, 'Conversiones a configurar', 2)
make_table(doc,
    ['Evento', 'Tipo', 'Valor asignado'],
    [
        ('Formulario completado (submit)', 'Conversión principal', 'Alto'),
        ('Click en WhatsApp', 'Conversión principal', 'Alto'),
        ('Click en teléfono', 'Conversión principal', 'Alto'),
        ('Tiempo en página > 3 min', 'Micro-conversión', 'Bajo'),
        ('Consulta alquiler (campo formulario KM)', 'Conversión segmentada', 'Alto'),
        ('Consulta compra (campo formulario KM)', 'Conversión segmentada', 'Medio'),
    ],
    col_widths=[6, 5, 5]
)

add_heading(doc, 'Presupuesto por campaña', 2)
make_table(doc,
    ['#', 'Campaña', 'CPC máx. (ARS)', 'CPC máx. (USD)', 'Diario (ARS)', 'Diario (USD)', 'Mensual (ARS)', 'Mensual (USD)', 'Prior.'],
    [
        ('01', 'KM Alquiler', '2.500', '2,1', '8.000', '6,7', '240.000', '200', 'P1'),
        ('02', 'KM Venta P/M', '2.000', '1,7', '8.000', '6,7', '240.000', '200', 'P1'),
        ('14', 'Insumos Gran Formato', '1.200', '1,0', '6.000', '5,0', '180.000', '150', 'P1'),
        ('03', 'KM Accurio', '3.000', '2,5', '5.000', '4,2', '150.000', '125', 'P2'),
        ('04', 'Insumos KM', '1.500', '1,25', '4.000', '3,3', '120.000', '100', 'P2'),
        ('15', 'Dlican UV Flatbed', '2.000', '1,7', '4.000', '3,3', '120.000', '100', 'P2'),
        ('16', 'Dlican UV Híbrida', '2.000', '1,7', '4.000', '3,3', '120.000', '100', 'P2'),
        ('17', 'Sinajet Mesas Corte', '1.800', '1,5', '4.000', '3,3', '120.000', '100', 'P2'),
        ('18', 'RISO', '1.800', '1,5', '4.000', '3,3', '120.000', '100', 'P2'),
        ('05', 'UV Cama Plana (HF-3040)', '1.000', '0,83', '2.500', '2,1', '75.000', '62', 'P3'),
        ('05a', 'HUENU HF-6090', '1.200', '1,0', '3.000', '2,5', '90.000', '75', 'P3'),
        ('06', 'UV Rollo (otros)', '1.000', '0,83', '2.500', '2,1', '75.000', '62', 'P3'),
        ('06a', 'HUENU HR1804UV', '1.200', '1,0', '3.500', '2,9', '105.000', '87', 'P3'),
        ('07', 'UV Híbrida', '1.000', '0,83', '2.500', '2,1', '75.000', '62', 'P3'),
        ('08', 'DTF Textil', '1.000', '0,83', '3.000', '2,5', '90.000', '75', 'P3'),
        ('09', 'DTF UV', '1.000', '0,83', '2.500', '2,1', '75.000', '62', 'P3'),
        ('10', 'Sublimación', '900', '0,75', '2.500', '2,1', '75.000', '62', 'P3'),
        ('11', 'Ecosolvente + Solvente', '900', '0,75', '2.500', '2,1', '75.000', '62', 'P3'),
        ('19', 'HUENU UV Rollo (interior)', '800', '0,67', '2.500', '2,1', '75.000', '62', 'P3'),
        ('20', 'HUENU UV Cama Plana (int.)', '700', '0,58', '2.000', '1,7', '60.000', '50', 'P3'),
        ('21', 'HUENU UV Híbrida (int.)', '700', '0,58', '2.000', '1,7', '60.000', '50', 'P3'),
        ('22', 'HUENU DTF Textil (int.)', '700', '0,58', '2.000', '1,7', '60.000', '50', 'P3'),
        ('23', 'HUENU DTF UV (int.)', '700', '0,58', '2.000', '1,7', '60.000', '50', 'P3'),
        ('24', 'HUENU Sublimación (int.)', '700', '0,58', '2.000', '1,7', '60.000', '50', 'P3'),
        ('12', 'Tenneth otros modelos', '700', '0,58', '2.000', '1,7', '60.000', '50', 'P4'),
        ('12a', 'Tenneth FC7090', '800', '0,67', '2.500', '2,1', '75.000', '62', 'P4'),
        ('13', 'Etiquetas', '800', '0,67', '2.000', '1,7', '60.000', '50', 'P4'),
        ('25', 'HUENU Distribuidores', '700', '0,58', '2.000', '1,7', '60.000', '50', 'P4'),
    ],
    col_widths=[0.8, 4.2, 2, 2, 2, 2, 2.5, 2.5, 1.2]
)

add_heading(doc, 'Presupuesto consolidado por fase de activación', 2)
make_table(doc,
    ['Fase', 'Campañas activas', 'Mensual (ARS)', 'Mensual (USD)', 'Momento'],
    [
        ('P1 — Lanzamiento', '01, 02, 14 (3 campañas)', '660.000', '550', 'Mes 1'),
        ('P2 — Escalada', '+ 03, 04, 15, 16, 17, 18 (6 campañas)', '+630.000', '+525', 'Mes 2'),
        ('ACUMULADO P1+P2', '9 campañas', '1.290.000', '1.075', '—'),
        ('P3 — Expansión', '+ 05, 05a, 06, 06a, 07-11, 19-24 (15 campañas)', '+1.005.000', '+838', 'Mes 3'),
        ('ACUMULADO P1+P2+P3', '24 campañas', '2.295.000', '1.913', '—'),
        ('P4 — Cobertura total', '+ 12, 12a, 13, 25 (4 campañas)', '+255.000', '+212', 'Mes 4–5'),
        ('TOTAL — 27 campañas', 'Todas activas', '2.550.000', '2.125', '—'),
    ],
    col_widths=[3.5, 5.5, 3, 3, 2]
)

add_heading(doc, 'CPA objetivo por tipo de lead', 2)
make_table(doc,
    ['Tipo de lead', 'CPA objetivo (ARS)', 'CPA objetivo (USD)'],
    [
        ('KM Alquiler — contrato', '60.000', '50'),
        ('KM Venta mediana', '40.000', '33'),
        ('KM Accurio', '80.000', '67'),
        ('Insumos KM', '8.000', '6,7'),
        ('Insumos Gran Formato', '6.000', '5'),
        ('UV / DTF / Sublimación AMBA', '25.000', '21'),
        ('Dlican / Sinajet / RISO', '80.000', '67'),
        ('Tenneth (corte entry)', '15.000', '12,5'),
        ('Etiquetas', '15.000', '12,5'),
        ('HUENU interior (usuario final)', '12.000', '10'),
        ('HUENU Distribuidores', '30.000', '25'),
    ],
    col_widths=[7, 4, 4]
)

add_heading(doc, 'Métricas de éxito por tipo de campaña', 2)
make_table(doc,
    ['Campaña', 'CTR objetivo', 'CPA objetivo', 'Métrica clave'],
    [
        ('KM Alquiler', '> 5%', 'ARS 60.000', 'Contratos firmados / leads'),
        ('KM Venta P/M', '> 4%', 'ARS 40.000', 'Leads calificados'),
        ('KM Accurio', '> 3%', 'ARS 80.000', 'Leads industriales'),
        ('Insumos (ambos)', '> 6%', 'ARS 6.000–8.000', 'Pedidos concretados'),
        ('UV / DTF / Sub / Eco AMBA', '> 4%', 'ARS 25.000', 'Leads con consulta de precio'),
        ('Dlican / Sinajet / RISO', '> 3%', 'ARS 80.000', 'Leads industriales calificados'),
        ('Tenneth / Etiquetas', '> 5%', 'ARS 15.000', 'Cotizaciones solicitadas'),
        ('HUENU interior', '> 3,5%', 'ARS 12.000', 'Consultas + impresiones de marca'),
        ('HUENU Distribuidores', '> 3%', 'ARS 30.000', 'Distribuidores incorporados'),
    ],
    col_widths=[4.5, 2.5, 3.5, 5.5]
)

page_break(doc)

add_heading(doc, 'Checklist de activación', 2)
add_heading(doc, 'Antes de activar cualquier campaña', 3)
pre = [
    'Conectar Google Ads con Google Analytics 4',
    'Configurar eventos de conversión: formulario submit, click WhatsApp, click teléfono',
    'Publicar las 4 landing pages en servidor con URL definitiva',
    'Verificar script WhatsApp en mobile y desktop (lp-km-alquiler, lp-hf6090, lp-hr1804uv, lp-teneth7090)',
    'Configurar gracias.html como confirmación de conversión',
    'Cargar logo, extensiones de imagen y assets en Google Ads',
    'Configurar exclusión geográfica AMBA para todas las campañas HUENU interior (19–25)',
    'Verificar que huenu.ar está online y el HR1804UV está publicado',
    'Crear landing de distribuidores en huenu.ar antes de activar campaña 25',
]
for item in pre:
    add_bullet(doc, '☐  ' + item)

add_heading(doc, 'P1 — Lanzamiento (Mes 1): Campañas 01, 02, 14', 3)
p1 = [
    '01 KM Alquiler — CPC máx. ARS 2.500 · Diario ARS 8.000 · LP publicada y WA testeado',
    '02 KM Venta P/M — CPC máx. ARS 2.000 · Diario ARS 8.000 · Avisos RSA cargados',
    '14 Insumos Gran Formato — CPC máx. ARS 1.200 · Diario ARS 6.000 · Keywords genéricas sin marca',
    'Configurar conversión segmentada en KM Alquiler: "consulta alquiler" vs "consulta compra"',
]
for item in p1:
    add_bullet(doc, '☐  ' + item)

add_heading(doc, 'P2 — Escalada (Mes 2): + Campañas 03, 04, 15, 16, 17, 18', 3)
p2 = [
    '03 KM Accurio — CPC ARS 3.000 · Diario ARS 5.000',
    '04 Insumos KM — CPC ARS 1.500 · Diario ARS 4.000',
    '15 Dlican Flatbed — CPC ARS 2.000 · Diario ARS 4.000',
    '16 Dlican Híbrida — CPC ARS 2.000 · Diario ARS 4.000',
    '17 Sinajet — CPC ARS 1.800 · Diario ARS 4.000',
    '18 RISO — CPC ARS 1.800 · Diario ARS 4.000',
    'Revisar informe de términos de búsqueda P1 y agregar negativas',
    'Ajustar presupuestos P1 según ROAS real antes de agregar P2',
]
for item in p2:
    add_bullet(doc, '☐  ' + item)

add_heading(doc, 'P3 — Expansión (Mes 3): + Campañas 05, 05a, 06, 06a, 07–11, 19–24', 3)
p3 = [
    'Activar campañas dedicadas 05a (HF-6090) y 06a (HR1804UV) antes que las genéricas 05 y 06',
    'Verificar LPs HF-6090 y HR1804UV con URL definitiva publicada',
    'Confirmar exclusión geográfica AMBA activa en campañas 19–24 (HUENU interior)',
    'Configurar seguimiento de impresiones y frecuencia para brand awareness HUENU',
    '"Pedísela a tu distribuidor" presente en al menos una descripción por campaña HUENU interior',
]
for item in p3:
    add_bullet(doc, '☐  ' + item)

add_heading(doc, 'P4 — Cobertura total (Mes 4–5): + Campañas 12, 12a, 13, 25', 3)
p4 = [
    '12 Tenneth otros modelos — CPC ARS 700 · Diario ARS 2.000',
    '12a Tenneth FC7090 — CPC ARS 800 · Diario ARS 2.500 · LP publicada',
    '13 Etiquetas — CPC ARS 800 · Diario ARS 2.000',
    '25 HUENU Distribuidores — CPC ARS 700 · Diario ARS 2.000 · Landing distribuidores en huenu.ar',
]
for item in p4:
    add_bullet(doc, '☐  ' + item)

# ── Guardar ────────────────────────────────────────────────────────────────────
output = r'C:\Users\jipi_\landing-pages\docs\Estrategia_GoogleAds_SistemasSoluciones_2026.docx'
doc.save(output)
print(f'OK: {output}')
