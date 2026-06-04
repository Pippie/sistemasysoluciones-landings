from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ORANGE = RGBColor(0xFF, 0x9B, 0x11)
DARK   = RGBColor(0x2E, 0x2E, 0x2E)
GRAY   = RGBColor(0x50, 0x50, 0x50)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

def set_cell_bg(cell, color_hex):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level <= 2 else 10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text); run.bold = True
    if level == 1:   run.font.size = Pt(16); run.font.color.rgb = ORANGE
    elif level == 2: run.font.size = Pt(13); run.font.color.rgb = DARK
    elif level == 3: run.font.size = Pt(11); run.font.color.rgb = GRAY
    else:            run.font.size = Pt(10); run.font.color.rgb = ORANGE

def add_body(doc, text, bold=False, italic=False, size=10):
    p = doc.add_paragraph(); run = p.add_run(text)
    run.bold = bold; run.italic = italic; run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(3)

def add_bullet(doc, text, size=9):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text); run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(2)

def make_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers)); t.style = 'Table Grid'
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        c = hrow.cells[i]; set_cell_bg(c, 'FF9B11')
        run = c.paragraphs[0].add_run(h); run.bold = True
        run.font.color.rgb = WHITE; run.font.size = Pt(9)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]
            run = c.paragraphs[0].add_run(str(val)); run.font.size = Pt(9)
            if ri % 2 == 0: set_cell_bg(c, 'F5F5F5')
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in t.rows: row.cells[ci].width = Cm(w)
    doc.add_paragraph(); return t

# ── Avisos faltantes ─────────────────────────────────────────────────────────
avisos = [
    {
        'num': '03', 'nombre': 'KM Accurio Alta Producción',
        'titulos': [
            'AccurioPress Konica Minolta', 'Impresión digital industrial', 'Offset digital de alta tirada',
            'AccurioPress C6100', '100 páginas por minuto', 'Konica Minolta Oficial',
            'Soporte técnico especializado', '40 años en Argentina', 'Cotizá tu AccurioPress',
            'Para imprentas industriales', 'Calidad offset digital', 'Alta tirada en color',
            'Consultá disponibilidad', 'Stock y financiación', 'Máximo uptime garantizado',
        ],
        'descripciones': [
            'AccurioPress C3070, C6100 y Bizhub Pro 1100. Calidad offset digital para alta tirada.',
            'Hasta 100 ppm en color. Ciclo mensual de 1.800.000 páginas. Distribuidores oficiales AMBA.',
            '40 años respaldando imprentas argentinas. Soporte técnico propio, sin tercerizar.',
            'Analizamos tu flujo de producción y te recomendamos el modelo exacto. Consultá hoy.',
        ],
        'url': 'sistemasysoluciones.com/Accurio-Produccion',
        'sitelinks': [
            ('AccurioPress C6100', '100 ppm color industrial', 'Ciclo 1.800.000 pág./mes'),
            ('AccurioPress C3070/80', '71–81 ppm alta producción', 'Encuadernación automática'),
            ('Bizhub Pro 1100', '100 ppm blanco y negro', 'Escáner dual simultáneo'),
            ('Pedí presupuesto', 'Sin compromiso ni costo', 'Asesor industrial en 24hs'),
        ],
        'callouts': ['Distribuidor oficial KM', 'Soporte técnico propio', '40 años en el mercado', 'Uptime garantizado'],
        'snippet': ('Modelos', 'AccurioPress C3070, AccurioPress C6100, Bizhub Pro 1100'),
    },
    {
        'num': '04', 'nombre': 'Insumos KM',
        'titulos': [
            'Tóners Konica Minolta', 'Repuestos Bizhub Originales', 'Insumos KM en Stock',
            'Tóner AccurioPress', 'Entrega en todo el país', 'Originales garantizados',
            '40 años en Argentina', 'Consultá stock hoy', 'Sin demoras en entrega',
            'Drum y repuestos Bizhub', 'Distribuidor oficial KM', 'Todo el país',
            'Pedí tu presupuesto', 'Calidad y rendimiento', 'Tóner Bizhub C658',
        ],
        'descripciones': [
            'Tóners y repuestos Konica Minolta 100% originales. Stock disponible. Entrega en todo el país.',
            'Drum, tóners y repuestos para Bizhub y AccurioPress. Rendimiento garantizado por fabricante.',
            '40 años distribuyendo insumos KM en Argentina. Soporte técnico y asesoramiento incluidos.',
            'Consultá disponibilidad y precio. Evitá fallas por insumos no originales. Pedí info hoy.',
        ],
        'url': 'sistemasysoluciones.com/Insumos-KM',
        'sitelinks': [
            ('Tóners Bizhub', 'Para modelos B/N y Color', '100% originales garantizados'),
            ('Tóners AccurioPress', 'Para alta producción', 'Máximo rendimiento por hoja'),
            ('Drums y repuestos', 'Piezas críticas en stock', 'Envío a todo el país'),
            ('Consultá tu modelo', 'Encontrá el insumo exacto', 'Asesor en 24hs hábiles'),
        ],
        'callouts': ['Originales garantizados', 'Stock permanente', 'Entrega todo el país', '40 años en el mercado'],
        'snippet': ('Tipos de insumo', 'Tóners, Drums, Repuestos, Kits de mantenimiento'),
    },
    {
        'num': '05', 'nombre': 'UV Cama Plana — HF-3040',
        'titulos': [
            'Impresora UV Cama Plana', 'HUENU HF-3040 Compacto', 'Entrada al UV sobre objetos',
            'Madera, vidrio y acrílico', 'UV LED con blanco y barniz', 'Soporte técnico local',
            '40 años en Argentina', 'Cotizá tu equipo hoy', 'Para taller pequeño',
            'Impresora UV A3 compacta', 'Reseller oficial HUENU', 'Capacitación incluida',
            'Repuestos en el país', 'Consultá precio y stock', 'Alta definición 1800 dpi',
        ],
        'descripciones': [
            'HUENU HF-3040: UV cama plana compacto, 297×420mm. Ideal para comenzar en impresión UV.',
            'Imprimí sobre madera, vidrio, acrílico y más. CMYK + Blanco + Barniz. UV LED.',
            'Reseller oficial HUENU en Argentina. Capacitación y repuestos incluidos. 40 años.',
            'La solución de entrada para personalización UV. Consultá precio y financiación hoy.',
        ],
        'url': 'sistemasysoluciones.com/productos/impresion/gran-formato/cama-plana',
        'sitelinks': [
            ('HUENU HF-3040', '297×420mm compacto', 'Ideal para empezar en UV'),
            ('HUENU HF-6090', '600×900mm mayor formato', 'Ver landing dedicada'),
            ('Materiales compatibles', 'Madera, vidrio, acrílico', 'Metal, cuero, packaging'),
            ('Pedí tu cotización', 'Sin compromiso ni costo', 'Respondemos en 24hs'),
        ],
        'callouts': ['Reseller oficial HUENU', 'Soporte técnico local', 'Capacitación incluida', 'Repuestos en stock'],
        'snippet': ('Aplicaciones', 'Merchandising, Packaging, Vidrio, Madera, Acrílico'),
    },
    {
        'num': '06', 'nombre': 'UV Rollo a Rollo — Otros modelos',
        'titulos': [
            'Plotter UV Rollo a Rollo', 'HUENU UV 1600 a 3200mm', 'Señalética UV de gran escala',
            'Hasta 12 cabezales Epson', 'Hasta 3200mm de ancho', 'Soporte técnico local',
            '40 años en Argentina', 'Cotizá tu plotter UV', 'Para señalética industrial',
            'UV con blanco y barniz', 'Reseller oficial HUENU', 'Capacitación incluida',
            'Alta producción UV', 'Consultá modelos y precio', 'Repuestos en el país',
        ],
        'descripciones': [
            'HUENU UV Rollo: desde HR1602UV hasta HR3212UV. Hasta 3200mm y 12 cabezales Epson.',
            'Señalética, viniles, wallpapers y comunicación visual a gran escala. CMYK+B+V+Flúo.',
            'Reseller oficial HUENU en Argentina. Capacitación, soporte técnico y repuestos locales.',
            '40 años en el mercado argentino. Consultá el modelo ideal para tu volumen de producción.',
        ],
        'url': 'sistemasysoluciones.com/productos/impresion/gran-formato/uv-rollo-a-rollo',
        'sitelinks': [
            ('HUENU HR2008UV', '2000mm, hasta 8 cabezales', 'Alta productividad industrial'),
            ('HUENU HR3208UV', '3200mm, 8 cabezales', 'Señalética de gran formato'),
            ('HUENU HR3212UV', '3200mm, 12 cabezales', 'Máxima producción UV'),
            ('Pedí tu cotización', 'Sin compromiso ni costo', 'Respondemos en 24hs'),
        ],
        'callouts': ['Hasta 3200mm de ancho', 'Reseller oficial HUENU', 'Soporte técnico local', 'Capacitación incluida'],
        'snippet': ('Modelos', 'HR1602UV, HR2008UV, HR3204UV, HR3208UV, HR3212UV'),
    },
    {
        'num': '07', 'nombre': 'UV Híbrida AMBA',
        'titulos': [
            'Impresora UV Híbrida', 'Rollo y rígidos en uno', 'HUENU HI1804UV Oficial',
            'Xenons R180 PRO', 'Máxima versatilidad UV', 'CMYK, Blanco y Barniz',
            'Soporte técnico local', '40 años en Argentina', 'Cotizá tu equipo hoy',
            'Para gráfica y señalética', '1800mm de ancho', 'Un equipo, dos formatos',
            'Consultá modelos y precio', 'Repuestos en el país', 'Capacitación incluida',
        ],
        'descripciones': [
            'HUENU HI1804UV y Xenons R180 PRO: híbridos 1800mm para rollo y rígidos en un equipo.',
            'No necesitás dos equipos. Señalética, packaging y comunicación visual en un solo plotter.',
            'Reseller oficial en Argentina. CMYK + Blanco + Barniz. UV LED. Capacitación incluida.',
            '40 años en el mercado argentino. Consultá el modelo ideal para tu tipo de producción.',
        ],
        'url': 'sistemasysoluciones.com/productos/impresion/gran-formato/uv-híbrido',
        'sitelinks': [
            ('HUENU HI1804UV', '1800mm híbrido, Epson I3200', 'Rígidos hasta 30mm altura'),
            ('Xenons R180 PRO', '1800mm, 4 cabezales Epson', 'CMYK+LC+LM+Blanco+Barniz'),
            ('Aplicaciones híbridas', 'Señalética, packaging', 'Vinilo y rígidos combinados'),
            ('Pedí tu cotización', 'Sin compromiso ni costo', 'Respondemos en 24hs'),
        ],
        'callouts': ['Rollo y rígidos en uno', 'Soporte técnico local', 'Capacitación incluida', '40 años en el mercado'],
        'snippet': ('Aplicaciones', 'Señalética, Packaging, Vinilo, Rígidos, Comunicación visual'),
    },
    {
        'num': '08', 'nombre': 'DTF Textil AMBA',
        'titulos': [
            'Plotter DTF Textil Precio', 'HUENU DTF Textil Oficial', 'Estampado DTF profesional',
            'Shaker industrial integrado', 'Epson I3200 hasta 48 m²/h', 'Soporte técnico local',
            '40 años en Argentina', 'Cotizá tu plotter DTF', 'Para remeras y textiles',
            'CMYK y Blanco garantizados', 'Reseller oficial HUENU', 'Capacitación incluida',
            'Consultá modelos y precio', 'Repuestos en el país', 'DTF de alta productividad',
        ],
        'descripciones': [
            'HUENU HR-DT series: DTF textil con shaker integrado. Desde entry level hasta 48 m²/h.',
            'Estampá remeras, telas y textiles con calidad profesional. CMYK + Blanco resistente al lavado.',
            'Reseller oficial HUENU en Argentina. Capacitación, soporte técnico y repuestos en el país.',
            '40 años en el mercado argentino. Consultá el modelo ideal para tu volumen de producción.',
        ],
        'url': 'sistemasysoluciones.com/productos/impresion/dtf/dtf-textil',
        'sitelinks': [
            ('HUENU HR6007DT', '620mm, 27–48 m²/h', 'Shaker y curado integrado'),
            ('HUENU HR6002DT', 'Entry level productivo', 'Ideal para empezar en DTF'),
            ('Hanrun Super A-602', '9 colores + fluorescentes', 'Alta productividad textil'),
            ('Pedí tu cotización', 'Sin compromiso ni costo', 'Respondemos en 24hs'),
        ],
        'callouts': ['Shaker industrial integrado', 'Reseller oficial HUENU', 'Soporte técnico local', 'Capacitación incluida'],
        'snippet': ('Modelos', 'HR3002DT, HR6002DT, HR6004DT, HR6007DT, Hanrun Super A-602'),
    },
    {
        'num': '09', 'nombre': 'DTF UV AMBA',
        'titulos': [
            'Impresora DTF UV Precio', 'HUENU DTF UV Oficial', 'DTF UV con laminado',
            'Para objetos y packaging', 'CMYK, Blanco y Barniz UV', 'Soporte técnico local',
            '40 años en Argentina', 'Cotizá tu equipo hoy', 'Transferencia sobre rígidos',
            'Curado UV LED integrado', 'Reseller oficial HUENU', 'Capacitación incluida',
            'Consultá modelos y precio', 'Repuestos en el país', 'Acabado premium garantizado',
        ],
        'descripciones': [
            'HUENU DTF UV: impresión y laminado integrado. 320mm y 620mm. CMYK + Blanco + Barniz UV.',
            'Transferencia de alta calidad sobre objetos, packaging y superficies rígidas o flexibles.',
            'Reseller oficial HUENU en Argentina. Capacitación incluida. Soporte técnico y repuestos locales.',
            '40 años en el mercado argentino. Consultá el modelo ideal para tu tipo de producción.',
        ],
        'url': 'sistemasysoluciones.com/productos/impresion/dtf/dtf-uv',
        'sitelinks': [
            ('HUENU HR6004DU', '620mm con laminado', 'Alta productividad DTF UV'),
            ('HUENU HR3001DU', '320mm compacto', 'Ideal para volumen moderado'),
            ('Hanrun UV-H6003S', 'Rollo a rollo DTF UV', 'Múltiples superficies'),
            ('Pedí tu cotización', 'Sin compromiso ni costo', 'Respondemos en 24hs'),
        ],
        'callouts': ['Laminado integrado', 'Reseller oficial HUENU', 'Soporte técnico local', 'Capacitación incluida'],
        'snippet': ('Aplicaciones', 'Packaging, Objetos rígidos, Merchandising, Etiquetas premium'),
    },
    {
        'num': '10', 'nombre': 'Sublimación Gran Formato AMBA',
        'titulos': [
            'Plotter Sublimación Precio', 'HUENU Sublimación Oficial', 'Impresión textil a escala',
            'Hasta 8 cabezales Epson', '300 m²/h de productividad', 'Soporte técnico local',
            '40 años en Argentina', 'Cotizá tu plotter hoy', 'Para telas y textiles',
            '1600mm a 2220mm de ancho', 'Reseller oficial HUENU', 'Capacitación incluida',
            'Consultá modelos y precio', 'Repuestos en el país', 'Alta velocidad de sublimación',
        ],
        'descripciones': [
            'HUENU HR-SU series: de 1600mm a 2220mm. Hasta 8 cabezales y 300 m²/h de productividad.',
            'Secado IR + aire de alta potencia. CMYK multicanal. Calidad y uniformidad en cada metro.',
            'Reseller oficial HUENU en Argentina. Capacitación, soporte técnico y repuestos incluidos.',
            '40 años en el mercado argentino. Encontrá el modelo ideal para tu volumen de sublimación.',
        ],
        'url': 'sistemasysoluciones.com/productos/impresion/gran-formato/sublimacion',
        'sitelinks': [
            ('HUENU HR1808SU', '8 cabezales, 300 m²/h', 'Ultra alta velocidad textil'),
            ('HUENU HR1804SU', '4 cabezales, 1800mm', 'Alta productividad'),
            ('HUENU HR1602SU', '2 cabezales, 1600mm', 'Eficiente y accesible'),
            ('Pedí tu cotización', 'Sin compromiso ni costo', 'Respondemos en 24hs'),
        ],
        'callouts': ['Hasta 300 m²/h', 'Reseller oficial HUENU', 'Soporte técnico local', 'Capacitación incluida'],
        'snippet': ('Modelos', 'HR1602SU, HR1804SU, HR1808SU, HR1902SU, HR2202SU'),
    },
    {
        'num': '11', 'nombre': 'Ecosolvente + Solvente Gran Formato',
        'titulos': [
            'Plotter Ecosolvente Precio', 'HUENU Ecosolvente Oficial', 'Señalética exterior durable',
            'Hasta 3200mm de ancho', 'Epson I3200, 60 m²/h', 'Soporte técnico local',
            '40 años en Argentina', 'Cotizá tu plotter hoy', 'Banners y comunicación',
            'Tintas resistentes al sol', 'Reseller oficial HUENU', 'Capacitación incluida',
            'Consultá modelos y precio', 'Repuestos en el país', 'Viniles para exteriores',
        ],
        'descripciones': [
            'HUENU ecosolvente: desde 1600mm hasta 3200mm. Hasta 60 m²/h. Tintas durables al exterior.',
            'Señalética, banners y viniles para exteriores con colores vibrantes y máxima durabilidad.',
            'Reseller oficial HUENU en Argentina. Capacitación incluida. Soporte técnico y repuestos locales.',
            '40 años en el mercado argentino. Consultá el modelo exacto para tu volumen de producción.',
        ],
        'url': 'sistemasysoluciones.com/productos/impresion/gran-formato/ecosolvente',
        'sitelinks': [
            ('HUENU HR3204ES', '3200mm, hasta 60 m²/h', 'Ultra gran formato exterior'),
            ('HUENU HR1804ES', '1800mm eficiente', 'Secado IR + aire de potencia'),
            ('HUENU HR1602ES', '1600mm accesible', 'Para comenzar en señalética'),
            ('Pedí tu cotización', 'Sin compromiso ni costo', 'Respondemos en 24hs'),
        ],
        'callouts': ['Tintas resistentes UV', 'Reseller oficial HUENU', 'Soporte técnico local', 'Capacitación incluida'],
        'snippet': ('Aplicaciones', 'Señalética exterior, Banners, Vinilos, Carteles, Comunicación visual'),
    },
    {
        'num': '12', 'nombre': 'Tenneth — Otros Modelos',
        'titulos': [
            'Mesas de Corte Tenneth', 'Corte digital de precisión', 'Tenneth FC9012 y FC1313U',
            'Camas hasta 1300×1300mm', 'Precisión ±0,1mm garantizada', 'Soporte técnico local',
            '40 años en Argentina', 'Cotizá tu mesa de corte', 'Con o sin transportadora',
            'Hasta 1000mm/s', 'Distribuidor oficial', 'Capacitación incluida',
            'Consultá modelos y precio', 'Corte y hendido en uno', 'Para mayor volumen de corte',
        ],
        'descripciones': [
            'Tenneth FC9012 (900×1200mm) y FC1313U (1300×1300mm con conveyor). Más capacidad de corte.',
            'Cámara CCD para corte de contorno. Precisión ±0,1mm. Hasta 1000mm/s. Doble herramienta.',
            'Distribuidor oficial Tenneth en Argentina. Capacitación incluida. Soporte técnico local.',
            '40 años en el mercado argentino. Automatizá tu producción y escalá sin más personal.',
        ],
        'url': 'sistemasysoluciones.com/productos/corte/mesas_de_corte/equipos-compatos-corte',
        'sitelinks': [
            ('Tenneth FC1313U', '1300×1300mm conveyor', 'Máximo volumen de corte'),
            ('Tenneth FC9012', '900×1200mm transportadora', 'Para producción media-alta'),
            ('Tenneth FC7090', '700×900mm — ver LP', 'Modelo más buscado de la línea'),
            ('Pedí tu cotización', 'Sin compromiso ni costo', 'Respondemos en 24hs'),
        ],
        'callouts': ['Cámara CCD de contorno', 'Precisión ±0,1mm', 'Soporte técnico local', 'Capacitación incluida'],
        'snippet': ('Modelos', 'FC5070E, FC7090A, FC7090U, FC9012, FC1313U'),
    },
    {
        'num': '13', 'nombre': 'Etiquetas — Impresión y Corte',
        'titulos': [
            'Impresora de Etiquetas', 'Cortadora de Etiquetas', 'Producción in-house',
            'Harpy y Bizpress Oficial', 'Duoblade Troqueladora', 'Soporte técnico local',
            '40 años en Argentina', 'Cotizá tu equipo hoy', 'Impresión rollo a rollo',
            'Etiquetas autoadhesivas', 'Eliminá la tercerización', 'Consultá modelos y precio',
            'Para packaging y logística', 'Capacitación incluida', 'Alta precisión de corte',
        ],
        'descripciones': [
            'Impresoras y cortadoras de etiquetas digitales. Harpy, Bizpress, Duoblade. CABA y GBA.',
            'Producción de etiquetas autoadhesivas in-house. Sin tercerización. Mayor flexibilidad.',
            'Distribuidor oficial en Argentina. Capacitación incluida. Soporte técnico local propio.',
            '40 años en el mercado argentino. Consultá el sistema ideal para tu volumen de etiquetas.',
        ],
        'url': 'sistemasysoluciones.com/productos/impresion/etiquetas',
        'sitelinks': [
            ('Harpy R440', 'Inkjet rollo, 100m/min', 'Etiquetas a alta velocidad'),
            ('Bizpress 13R', 'Láser rollo, datos variables', 'PET, PP, PVC y más'),
            ('Duoblade WS Max', 'Cortadora 2 cabezales', 'Precisión 0,1mm a 9m/min'),
            ('Pedí tu cotización', 'Sin compromiso ni costo', 'Respondemos en 24hs'),
        ],
        'callouts': ['Producción in-house', 'Soporte técnico local', '40 años en el mercado', 'Capacitación incluida'],
        'snippet': ('Equipos', 'Harpy R440, Bizpress 13R, Duoblade WS Max, Duoblade SX, Teneth RN3'),
    },
    {
        'num': '14', 'nombre': 'Insumos Gran Formato',
        'titulos': [
            'Tinta UV para Plotter', 'Film DTF Textil', 'Polvo DTF Adhesivo',
            'Cabezal Epson I3200', 'Insumos para cualquier marca', 'Stock permanente nacional',
            '40 años en Argentina', 'Consultá precio y stock', 'Tinta sublimación Argentina',
            'Tinta ecosolvente compatible', 'Entrega en todo el país', 'Cabezal Ricoh repuesto',
            'Pedí tu presupuesto hoy', 'Calidad garantizada', 'Film DTF mate y brillante',
        ],
        'descripciones': [
            'Tintas UV, sublimación, ecosolvente y DTF. Films y polvos. Para cualquier marca de plotter.',
            'Cabezales Epson I3200 y Ricoh. Film DTF mate, brillante, hot y cold peel. Stock permanente.',
            '40 años distribuyendo insumos en Argentina. Calidad garantizada. Entrega en todo el país.',
            'Consultá precio, disponibilidad y compatibilidad con tu equipo. Respondemos en 24hs.',
        ],
        'url': 'sistemasysoluciones.com/productos/insumos-y-repuestos/tintas',
        'sitelinks': [
            ('Tintas UV y ecosolvente', 'Compatibles con cualquier marca', 'Calidad garantizada'),
            ('Films DTF', 'Mate, brillante, hot peel', 'Cold peel disponible'),
            ('Polvos DTF', 'Fino, grueso, antiamarillamiento', 'Alta adherencia textil'),
            ('Cabezales y repuestos', 'Epson I3200 y Ricoh', 'Envío a todo el país'),
        ],
        'callouts': ['Stock permanente', 'Para cualquier marca', 'Entrega todo el país', '40 años en el mercado'],
        'snippet': ('Tipos', 'Tintas UV, Tintas sublimación, Films DTF, Polvos DTF, Cabezales'),
    },
    {
        'num': '15', 'nombre': 'Dlican UV Cama Plana Alta Producción',
        'titulos': [
            'Dlican UV Flatbed Industrial', 'UV Cama Plana Alta Producción', 'Hasta 3300×2500mm de cama',
            'Ricoh G5/G6 industrial', '50 a 75 m²/h de producción', 'Para señalética industrial',
            '40 años en Argentina', 'Cotizá tu Dlican hoy', 'Posicionamiento visual 4K',
            'Distribuidor oficial Dlican', 'Todo Argentina', 'Soporte técnico propio',
            'Consultá modelos y precio', 'CMYK, Blanco y Barniz', 'Alta productividad UV',
        ],
        'descripciones': [
            'Dlican UV flatbed: desde 900×600 hasta 3300×2500mm. Cabezales Ricoh G5/G6. Alta producción.',
            'Posicionamiento automático 4K, conveyor y alturas hasta 350mm. Para señalética industrial.',
            'Distribuidor oficial Dlican en todo Argentina. 40 años de trayectoria. Soporte propio.',
            'Inversión con respaldo real. Consultá el modelo exacto para tu volumen industrial.',
        ],
        'url': 'sistemasysoluciones.com/productos/impresion/gran-formato/cama-plana',
        'sitelinks': [
            ('Dlican DLI-3325', '3300×2500mm industrial', 'Ricoh G5/G6, 50–75 m²/h'),
            ('Dlican DLI-2513', '2500×1300mm + pos. 4K', 'Máxima precisión automática'),
            ('Dlican DLI-1010', '1000×1000mm inteligente', 'Pos. automático visual'),
            ('Pedí tu cotización', 'Asesor industrial disponible', 'Respondemos en 24hs'),
        ],
        'callouts': ['Distribuidor oficial Dlican', 'Soporte técnico propio', 'Todo Argentina', '40 años en el mercado'],
        'snippet': ('Modelos', 'DLI-9060, DLI-1010, DLI-1215, DLI-1612, DLI-2513, DLI-3220, DLI-3325'),
    },
    {
        'num': '16', 'nombre': 'Dlican UV Híbrida Alta Producción',
        'titulos': [
            'Dlican UV Híbrida Industrial', 'Rollo y rígidos, alta producción', 'Hasta 6600mm de ancho',
            'Ricoh industrial 7PL', 'Hasta 430 m²/h', 'Para señalética a gran escala',
            '40 años en Argentina', 'Cotizá tu Dlican híbrido', 'Distribuidor oficial Dlican',
            'Todo Argentina', 'Soporte técnico propio', 'CMYK, Blanco y Barniz',
            'Consultá modelos y precio', 'Máxima productividad UV', 'Formato ultra gran escala',
        ],
        'descripciones': [
            'Dlican UV híbrida: de 1640mm a 6600mm. Ricoh Gen5/Gen6. Hasta 430 m²/h de producción.',
            'Un solo equipo para rollo y rígidos a escala industrial. Señalética, viniles y más.',
            'Distribuidor oficial Dlican en todo Argentina. 40 años de trayectoria. Soporte propio.',
            'Inversión industrial con respaldo real. Consultá el modelo ideal para tu operación.',
        ],
        'url': 'sistemasysoluciones.com/productos/impresion/gran-formato/uv-h%C3%ADbrido',
        'sitelinks': [
            ('Dlican DLI-6600', '6600mm, 236–400 m²/h', 'El más grande de la línea'),
            ('Dlican DLI-3300s', '3250mm, Ricoh 7PL', 'Hasta 430 m²/h industrial'),
            ('Dlican DLI-1688', '1640mm versátil', 'Entrada a híbridos Dlican'),
            ('Pedí tu cotización', 'Asesor industrial disponible', 'Respondemos en 24hs'),
        ],
        'callouts': ['Hasta 6600mm de ancho', 'Distribuidor oficial Dlican', 'Soporte técnico propio', 'Todo Argentina'],
        'snippet': ('Modelos', 'DLI-1688, DLI-2500, DLI-3300, DLI-3300s, DLI-5250, DLI-6600'),
    },
    {
        'num': '17', 'nombre': 'Sinajet Mesas de Corte Alta Producción',
        'titulos': [
            'Mesa de Corte Sinajet', 'Corte digital industrial', 'Hasta 1700mm/s de velocidad',
            'Precisión ≤0,1mm', 'Series DG, DH y DF', 'Para packaging industrial',
            '40 años en Argentina', 'Cotizá tu Sinajet hoy', 'Sin troquel, cero desperdicio',
            'Distribuidor oficial', 'Todo Argentina', 'Soporte técnico propio',
            'Consultá modelos y precio', 'Alta producción sin pausa', 'Multi-herramienta digital',
        ],
        'descripciones': [
            'Sinajet series DG, DH y DF. Hasta 1700mm/s, 45mm de espesor, precisión ≤0,1mm.',
            'Corte digital sin troquel para gráfica, packaging y textil. Mesa fija y con conveyor.',
            'Distribuidor oficial Sinajet en todo Argentina. 40 años de trayectoria. Soporte propio.',
            'Eliminá el corte manual. Escalá producción sin escalar personal. Consultá hoy.',
        ],
        'url': 'sistemasysoluciones.com/productos/corte/mesas_de_corte/serie-dg',
        'sitelinks': [
            ('Sinajet Serie DH', 'Alta producción industrial', '1700mm/s, hasta 45mm espesor'),
            ('Sinajet Serie DG', 'Producción versátil', 'Multi-herramienta intercambiable'),
            ('Sinajet Serie DF-MT', 'Conveyor, 1400mm/s', 'Compacta y eficiente'),
            ('Pedí tu cotización', 'Asesor industrial disponible', 'Respondemos en 24hs'),
        ],
        'callouts': ['Sin troquel', 'Precisión ≤0,1mm', 'Distribuidor oficial Sinajet', 'Todo Argentina'],
        'snippet': ('Series', 'Serie DG, Serie DF-MT, Serie DH, Serie DF'),
    },
    {
        'num': '18', 'nombre': 'RISO Inkjet Alta Velocidad',
        'titulos': [
            'Impresora RISO ComColor', 'Hasta 130 páginas por minuto', 'Inkjet alta velocidad',
            'Para mailings y tiradas', 'RISO GD7330 y FW5230', 'Bajo costo por página',
            '40 años en Argentina', 'Cotizá tu RISO hoy', 'Distribuidor oficial RISO',
            'Todo Argentina', 'Soporte técnico propio', 'Inkjet línea industrial',
            'Consultá modelos y precio', 'Alta tirada eficiente', '600×600 dpi en línea',
        ],
        'descripciones': [
            'RISO ComColor GD7330: 130 ppm. FW5230: 120 ppm dúplex. Inkjet de línea para alta tirada.',
            'El más bajo costo por página en impresión digital de alta velocidad. Mailing y producción.',
            'Distribuidor oficial RISO en todo Argentina. 40 años de trayectoria. Soporte técnico propio.',
            'Reemplazá el offset con tecnología inkjet de línea. Consultá el modelo para tu volumen.',
        ],
        'url': 'sistemasysoluciones.com/productos/impresion/pliegos/inkjet',
        'sitelinks': [
            ('RISO ComColor GD7330', '130 ppm, 600×600 dpi', 'Bandeja hasta 3500 hojas'),
            ('RISO ComColor FW5230', '120 ppm simplex', 'Dúplex 60 hojas/min'),
            ('Para mailing y producción', 'Bajo costo por hoja', 'Tiradas largas sin pausa'),
            ('Pedí tu cotización', 'Asesor industrial disponible', 'Respondemos en 24hs'),
        ],
        'callouts': ['Hasta 130 ppm', 'Bajo costo por página', 'Distribuidor oficial RISO', 'Todo Argentina'],
        'snippet': ('Modelos', 'ComColor GD7330, ComColor FW5230'),
    },
    {
        'num': '20', 'nombre': 'HUENU UV Cama Plana — Interior',
        'titulos': [
            'Impresora UV HUENU', 'UV Cama Plana HUENU', 'HUENU HF-6090',
            'Imprimí sobre objetos', 'Madera, vidrio y acrílico', 'CMYK, Blanco y Barniz UV',
            'Soporte local en Argentina', 'Pedísela a tu distribuidor', 'Para merchandising y deco',
            '600×900mm de cama', 'Marca HUENU en Argentina', 'Red de distribuidores',
            'Consultá donde conseguirla', 'Alta definición 1800 dpi', 'Calidad con respaldo local',
        ],
        'descripciones': [
            'HUENU HF-6090: impresión UV directa sobre cualquier objeto rígido. 600×900mm, 1800 dpi.',
            'Madera, vidrio, acrílico, metal, packaging. CMYK+Blanco+Barniz. Medición automática.',
            'HUENU tiene distribuidores en todo el interior del país. Pedísela a tu proveedor local.',
            'Soporte y repuestos disponibles en Argentina. Consultá dónde conseguir tu HUENU hoy.',
        ],
        'url': 'huenu.ar/UV-Plana',
        'sitelinks': [
            ('HUENU HF-6090', '600×900mm, 1800 dpi', 'Hasta 90mm de altura'),
            ('HUENU HF-3040', '297×420mm compacto', 'Ideal para empezar'),
            ('Materiales compatibles', 'Vidrio, madera, acrílico', 'Packaging y merchandising'),
            ('Red de distribuidores', 'Encontrá tu distribuidor', 'En todo el interior'),
        ],
        'callouts': ['Marca con respaldo local', 'Red de distribuidores', 'Repuestos en Argentina', 'Soporte técnico'],
        'snippet': ('Aplicaciones', 'Merchandising, Packaging, Vidrio, Madera, Acrílico, Metal'),
    },
    {
        'num': '21', 'nombre': 'HUENU UV Híbrida — Interior',
        'titulos': [
            'Plotter UV Híbrido HUENU', 'Rollo y rígidos HUENU', 'HUENU HI1804UV',
            'Un equipo, dos formatos', 'CMYK, Blanco y Barniz', 'Epson I3200 UV LED',
            'Soporte local en Argentina', 'Pedísela a tu distribuidor', 'Para gráfica y señalética',
            '1800mm de ancho', 'Marca HUENU en Argentina', 'Red de distribuidores',
            'Consultá donde conseguirla', 'Versatilidad UV profesional', 'Calidad con respaldo local',
        ],
        'descripciones': [
            'HUENU HI1804UV: híbrido 1800mm. Rollo y rígidos hasta 30mm. Epson I3200, UV LED.',
            'No necesitás dos equipos. Un solo plotter para señalética en vinilo y sobre rígidos.',
            'HUENU tiene distribuidores en todo el interior del país. Pedísela a tu proveedor local.',
            'Soporte y repuestos disponibles en Argentina. Consultá dónde conseguir tu HUENU hoy.',
        ],
        'url': 'huenu.ar/UV-Hibrida',
        'sitelinks': [
            ('HUENU HI1804UV', '1800mm híbrido', 'Rígidos hasta 30mm'),
            ('Rollo y rígidos', 'Vinilo y señalética', 'Sin cambiar de equipo'),
            ('Red de distribuidores', 'Encontrá tu distribuidor', 'En todo el interior'),
            ('Consultá tu modelo', 'Sin compromiso ni costo', 'Respondemos en 24hs'),
        ],
        'callouts': ['Rollo y rígidos en uno', 'Marca con respaldo local', 'Red de distribuidores', 'Soporte técnico'],
        'snippet': ('Aplicaciones', 'Señalética, Vinilo, Rígidos, Packaging, Comunicación visual'),
    },
    {
        'num': '22', 'nombre': 'HUENU DTF Textil — Interior',
        'titulos': [
            'Plotter DTF Textil HUENU', 'DTF para remeras HUENU', 'HUENU HR6007DT',
            'Shaker industrial integrado', 'Hasta 48 m²/h de producción', 'CMYK y Blanco garantizados',
            'Soporte local en Argentina', 'Pedísela a tu distribuidor', 'Para taller textil propio',
            '620mm de ancho de rollo', 'Marca HUENU en Argentina', 'Red de distribuidores',
            'Consultá donde conseguirla', 'Empezá con DTF profesional', 'Calidad con respaldo local',
        ],
        'descripciones': [
            'HUENU DTF textil: desde entry level hasta 48 m²/h. Shaker integrado. Cabezal Epson I3200.',
            'Estampá remeras, telas y textiles con calidad profesional. Resistente al lavado y duradero.',
            'HUENU tiene distribuidores en todo el interior del país. Pedísela a tu proveedor local.',
            'Soporte y repuestos disponibles en Argentina. Consultá dónde conseguir tu HUENU hoy.',
        ],
        'url': 'huenu.ar/DTF-Textil',
        'sitelinks': [
            ('HUENU HR6007DT', '620mm, 27–48 m²/h', 'Máxima productividad textil'),
            ('HUENU HR6002DT', 'Entry level productivo', 'Empezá en DTF textil'),
            ('HUENU HR3002DT', '330mm compacto', 'Ideal para pequeño taller'),
            ('Red de distribuidores', 'Encontrá tu distribuidor', 'En todo el interior'),
        ],
        'callouts': ['Shaker integrado', 'Marca con respaldo local', 'Red de distribuidores', 'Soporte técnico'],
        'snippet': ('Modelos', 'HR3002DT, HR6002DT, HR6004DT, HR6007DT'),
    },
    {
        'num': '23', 'nombre': 'HUENU DTF UV — Interior',
        'titulos': [
            'Plotter DTF UV HUENU', 'DTF UV con laminado', 'HUENU HR6004DU',
            'Para objetos y packaging', 'CMYK, Blanco y Barniz UV', 'Curado UV LED integrado',
            'Soporte local en Argentina', 'Pedísela a tu distribuidor', 'Transferencia sobre rígidos',
            '620mm de ancho', 'Marca HUENU en Argentina', 'Red de distribuidores',
            'Consultá donde conseguirla', 'Acabado premium en objetos', 'Calidad con respaldo local',
        ],
        'descripciones': [
            'HUENU DTF UV: impresión + laminado integrado. 320mm y 620mm. CMYK+Blanco+Barniz UV.',
            'Transferencia de alta calidad sobre objetos, packaging y superficies rígidas o flexibles.',
            'HUENU tiene distribuidores en todo el interior del país. Pedísela a tu proveedor local.',
            'Soporte y repuestos disponibles en Argentina. Consultá dónde conseguir tu HUENU hoy.',
        ],
        'url': 'huenu.ar/DTF-UV',
        'sitelinks': [
            ('HUENU HR6004DU', '620mm con laminado', 'Alta productividad DTF UV'),
            ('HUENU HR3001DU', '320mm compacto', 'Ideal para volumen moderado'),
            ('Aplicaciones DTF UV', 'Packaging, objetos', 'Superficies rígidas y flex.'),
            ('Red de distribuidores', 'Encontrá tu distribuidor', 'En todo el interior'),
        ],
        'callouts': ['Laminado integrado', 'Marca con respaldo local', 'Red de distribuidores', 'Soporte técnico'],
        'snippet': ('Aplicaciones', 'Packaging, Objetos rígidos, Merchandising, Etiquetas premium'),
    },
    {
        'num': '24', 'nombre': 'HUENU Sublimación Textil — Interior',
        'titulos': [
            'Plotter Sublimación HUENU', 'Sublimación textil HUENU', 'HUENU HR1808SU',
            'Hasta 300 m²/h', '8 cabezales Epson I3200', 'Uniformidad de color total',
            'Soporte local en Argentina', 'Pedísela a tu distribuidor', 'Para telas y textiles',
            'De 1600 a 2220mm de ancho', 'Marca HUENU en Argentina', 'Red de distribuidores',
            'Consultá donde conseguirla', 'Sublimación de alta escala', 'Calidad con respaldo local',
        ],
        'descripciones': [
            'HUENU sublimación: de 1600mm a 2220mm. Hasta 8 cabezales y 300 m²/h de producción.',
            'Secado IR + aire de alta potencia. CMYK multicanal. Calidad uniforme en toda la tirada.',
            'HUENU tiene distribuidores en todo el interior del país. Pedísela a tu proveedor local.',
            'Soporte y repuestos disponibles en Argentina. Consultá dónde conseguir tu HUENU hoy.',
        ],
        'url': 'huenu.ar/Sublimacion',
        'sitelinks': [
            ('HUENU HR1808SU', '8 cabezales, 300 m²/h', 'Ultra alta velocidad textil'),
            ('HUENU HR1804SU', '4 cabezales productivo', 'Alta calidad y velocidad'),
            ('HUENU HR1602SU', '1600mm eficiente', 'Ideal para empezar'),
            ('Red de distribuidores', 'Encontrá tu distribuidor', 'En todo el interior'),
        ],
        'callouts': ['Hasta 300 m²/h', 'Marca con respaldo local', 'Red de distribuidores', 'Soporte técnico'],
        'snippet': ('Modelos', 'HR1602SU, HR1804SU, HR1808SU, HR1902SU, HR2202SU'),
    },
]

# ── Abrir el documento existente y agregar sección ────────────────────────────
path = r'C:\Users\jipi_\landing-pages\docs\Estrategia_GoogleAds_SistemasSoluciones_2026.docx'
doc  = Document(path)

doc.add_page_break()
add_heading(doc, 'FASE 3 — AVISOS RSA COMPLETOS (campañas restantes)', 1)
add_body(doc, 'Avisos RSA para las 20 campañas no incluidas en la sección anterior. Mismo formato: 15 títulos + 4 descripciones + extensiones.', italic=True)

for av in avisos:
    doc.add_paragraph()
    add_heading(doc, f"Campaña {av['num']} · {av['nombre']}", 2)
    add_heading(doc, '15 Títulos (≤30 caracteres cada uno)', 3)
    for i, t in enumerate(av['titulos'], 1):
        add_bullet(doc, f"{i:02d}. {t}")
    add_heading(doc, '4 Descripciones (≤90 caracteres cada una)', 3)
    for i, d in enumerate(av['descripciones'], 1):
        add_bullet(doc, f"D{i}: {d}")
    add_body(doc, f"URL visible: {av['url']}", bold=True, size=9)
    add_heading(doc, 'Extensiones de sitio', 3)
    make_table(doc,
        ['Texto del sitelink', 'Descripción línea 1', 'Descripción línea 2'],
        av['sitelinks'],
        col_widths=[4.5, 5, 5]
    )
    add_body(doc, 'Extensión de llamada: (011) 4342-5742', size=9)
    add_body(doc, 'Textos destacados: ' + ' · '.join(av['callouts']), size=9)
    add_body(doc, f"Snippet estructurado — {av['snippet'][0]}: {av['snippet'][1]}", size=9)

doc.save(path)
print(f'OK: {len(avisos)} campanas agregadas: {path}')
